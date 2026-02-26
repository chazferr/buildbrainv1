"""
BuildBrain Extraction Engine
=============================
Shared extraction logic used by both the CLI (extract_to_excel.py)
and the web app (app.py).
"""

import base64
import csv
import email
import io
import json
import pathlib
import re
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from email import policy
from pathlib import Path
from typing import Callable, Optional

import anthropic
import fitz  # PyMuPDF
import google.generativeai as genai
import pandas as pd
from pydantic import BaseModel, ValidationError, field_validator

# ─── Gemini Vision ───────────────────────────────────────────────────────────
GEMINI_MODEL = "gemini-3.1-pro-preview"
_gemini_configured = False


def _get_gemini_client(api_key: str):
    global _gemini_configured
    if not _gemini_configured:
        genai.configure(api_key=api_key)
        _gemini_configured = True
    return genai.GenerativeModel(model_name=GEMINI_MODEL)

from wage_rates import get_wage, CT_WAGE_RATES
from material_db import get_material_price, MATERIAL_DB
from productivity_rates import calculate_labor, PRODUCTIVITY_RATES

# Optional imports for extended file types
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# ─── Rate Table Loader ────────────────────────────────────────────────────────

def _load_rate_tables():
    """Load rate tables from rate_tables.json. Returns dict or empty dict on failure."""
    rate_file = pathlib.Path(__file__).parent / "rate_tables.json"
    if rate_file.exists():
        try:
            with open(rate_file, encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"[WARNING] Could not load rate_tables.json: {e}")
    return {}


def _match_rate_profile(rate_tables: dict, project_quantities: dict) -> dict:
    """Find the best matching rate profile for this project.
    Returns the quotes dict from the best match, or empty dict.

    All project types (including single_family) go through the normal
    scoring system to find the best matching profile.
    """
    if not rate_tables or 'profiles' not in rate_tables:
        return {}

    pq = project_quantities or {}
    pt = pq.get('project_type', 'single_family')

    # Find best matching profile for any project type
    floors = pq.get('floor_count', 1) or 1
    units = pq.get('unit_count') or 0
    total_sf = pq.get('total_building_sf', 0) or 0

    best_match = None
    best_score = -1

    for profile in rate_tables['profiles']:
        chars = profile.get('characteristics', {})
        score = 0

        # Project type match (bonus, not required -- any commercial profile
        # is better than falling back to material_db)
        if pt in chars.get('project_type', []):
            score += 10
        else:
            score += 1  # still usable, just not ideal match

        # Floor count in range
        if chars.get('floor_count_min', 0) <= floors <= chars.get('floor_count_max', 999):
            score += 5

        # Unit count in range
        if chars.get('unit_count_min', 0) <= units <= chars.get('unit_count_max', 9999):
            score += 5

        # SF in range (don't gate on this -- SF extraction is unreliable)
        if total_sf > 0 and chars.get('total_sf_min', 0) <= total_sf <= chars.get('total_sf_max', 99999999):
            score += 3

        if score > best_score:
            best_score = score
            best_match = profile

    if best_match:
        pid = best_match.get('id', 'unknown')
        print(f"[RATE TABLE] project_type={pt} -> matched profile '{pid}' (score={best_score})")
        print(f"[PRICING SOURCE] project_type='{pt}' -> rate_tables")
        return best_match.get('quotes', {})

    print(f"[RATE TABLE] project_type={pt} -> no profiles in rate_tables.json, falling back to material_db")
    print(f"[PRICING SOURCE] project_type='{pt}' -> material_db")
    return {}


# Module-level load (once at import time)
_RATE_TABLES = _load_rate_tables()

# Mapping from TRADE_MAP trade names to rate table quote keys
_TRADE_TO_RATE_KEY = {
    "Building Concrete": "concrete_foundation",
    "Fire Sprinkler": "fire_sprinkler",
    "Roofing": "roofing_epdm",
    "HVAC": "hvac",
    "Electrical": "electrical",
    "Plumbing": "plumbing",
    "Drywall": "drywall",
    "Insulation": "insulation",
    "Painting": "painting",
    "Flooring": "flooring",
    "Windows": "windows_material",
    "Doors/Hdwr/Finish Carp": "doors_hardware_material",
    "Cabinets": "cabinets_install",
    "Specialties": "bathroom_accessories",
    "Siding": "siding_vinyl",
    "Sitework": "sitework",
    "SITE WORK / CIVIL": "sitework",
    "Rough Carpentry": "rough_carpentry_combined",
    "Masonry": "masonry",
}


def _apply_rate_table_override(trade_name, rate_profile, project_quantities):
    """Apply rate table pricing for a trade if a profile matches.

    Returns (budget_val, note_val) if rate table applies, or (None, None) if not.
    """
    if not rate_profile:
        return None, None

    rt_key = _TRADE_TO_RATE_KEY.get(trade_name)
    if not rt_key:
        return None, None

    # Special case: Rough Carpentry combines materials + labor rate table entries
    if rt_key == 'rough_carpentry_combined':
        mat_rt = rate_profile.get('rough_carpentry_materials')
        lab_rt = rate_profile.get('rough_carpentry_labor')
        if not mat_rt and not lab_rt:
            return None, None

        pq = project_quantities or {}
        total_sf = pq.get('total_building_sf') or 1000

        mat_rate = (mat_rt or {}).get('recommended_rate', 0) or 0
        lab_rate = (lab_rt or {}).get('recommended_rate', 0) or 0
        combined_rate = mat_rate + lab_rate

        if combined_rate <= 0:
            return None, None

        budget_val = round(combined_rate * total_sf)
        mat_bids = (mat_rt or {}).get('bids', [])
        lab_bids = (lab_rt or {}).get('bids', [])
        mat_note = (mat_rt or {}).get('notes', '')
        lab_note = (lab_rt or {}).get('notes', '')

        note_val = (
            f"Material ${mat_rate:.2f}/SF + Labor ${lab_rate:.2f}/SF "
            f"= ${combined_rate:.2f}/SF \u00d7 {total_sf:,.0f} SF | "
            f"Material: {mat_bids[0]['bidder'] if mat_bids else 'N/A'} "
            f"@ ${mat_bids[0].get('amount', 0):,.0f}. "
            f"Labor: {lab_bids[0]['bidder'] if lab_bids else 'N/A'} "
            f"@ ${lab_bids[0].get('amount', 0):,.0f}. | "
            f"{mat_note} {lab_note}"
        )
        return budget_val, note_val

    if rt_key not in rate_profile:
        return None, None

    rt = rate_profile[rt_key]
    rec_rate = rt.get('recommended_rate')
    rt_unit = rt.get('unit', 'per_total_sf')

    if rec_rate is not None and rec_rate > 0:
        pq = project_quantities or {}
        total_sf = pq.get('total_building_sf') or 1000
        footprint_sf = pq.get('footprint_sf') or (total_sf / max(pq.get('floor_count', 1) or 1, 1))
        unit_count = pq.get('unit_count') or 1
        plumbing_fixtures = pq.get('plumbing_fixtures') or rt.get('fixture_count_basis') or (unit_count * 5)
        door_count = (pq.get('ext_door_count') or max(2, unit_count // 10)) + (pq.get('int_door_count') or unit_count * 4)
        window_count = pq.get('window_count') or rt.get('opening_count_basis') or (unit_count * 3)

        if rt_unit == 'per_total_sf':
            budget_val = round(rec_rate * total_sf)
            qty_label = f"${rec_rate:,.2f}/SF \u00d7 {total_sf:,.0f} SF"
        elif rt_unit == 'per_footprint_sf':
            budget_val = round(rec_rate * footprint_sf)
            qty_label = f"${rec_rate:,.2f}/SF \u00d7 {footprint_sf:,.0f} footprint SF"
        elif rt_unit == 'per_fixture':
            budget_val = round(rec_rate * plumbing_fixtures)
            qty_label = f"${rec_rate:,.0f}/fixture \u00d7 {plumbing_fixtures} fixtures"
        elif rt_unit == 'per_opening':
            count = window_count if 'window' in rt_key else door_count
            budget_val = round(rec_rate * count)
            qty_label = f"${rec_rate:,.0f}/opening \u00d7 {count} openings"
        elif rt_unit == 'per_unit':
            budget_val = round(rec_rate * unit_count)
            qty_label = f"${rec_rate:,.0f}/unit \u00d7 {unit_count} units"
        elif rt_unit == 'lump_sum':
            bids = rt.get('bids', [])
            if bids:
                budget_val = bids[0].get('amount', round(rec_rate))
            else:
                budget_val = round(rec_rate)
            qty_label = f"lump sum ${budget_val:,.0f}"
        else:
            budget_val = round(rec_rate * total_sf)
            qty_label = f"${rec_rate:,.2f}/SF \u00d7 {total_sf:,.0f} SF"

        # Build rich note citing bidders
        bids = rt.get('bids', [])
        scope_note = rt.get('notes', '')
        bid_count = len(bids)
        if bid_count > 0:
            low_bidder = bids[0].get('bidder', 'Unknown')
            low_amount = bids[0].get('amount', 0)
            range_low = rt.get('range_low', rec_rate)
            range_high = rt.get('range_high', rec_rate)
            note_val = (
                f"{scope_note} | {qty_label} | "
                f"Based on {bid_count} actual bids "
                f"(low: {low_bidder} @ ${low_amount:,.0f}). "
                f"Range: ${range_low:,.2f}-${range_high:,.2f}/{rt_unit.replace('per_', '').replace('_', ' ')}."
            )
        else:
            note_val = f"{scope_note} | {qty_label} | Rate from Lazarus project reference data."

        return budget_val, note_val

    elif rt_unit == 'lump_sum':
        # lump_sum with no recommended_rate — use low bid amount
        bids = rt.get('bids', [])
        if bids and bids[0].get('amount'):
            budget_val = bids[0]['amount']
            scope_note = rt.get('notes', '')
            low_bidder = bids[0].get('bidder', 'Unknown')
            note_val = (
                f"{scope_note} | lump sum ${budget_val:,.0f} | "
                f"Low bid: {low_bidder} @ ${budget_val:,.0f}."
            )
            return budget_val, note_val

    return None, None


# ─── Constants ───────────────────────────────────────────────────────────────

MODEL = "claude-sonnet-4-5-20250929"
TEMPERATURE = 0
DPI = 200
MAX_IMAGE_BYTES = 4_800_000  # Stay under Claude's 5MB base64 limit

# Cost rates (USD per 1M tokens)
INPUT_RATE = 18.0
OUTPUT_RATE = 85.0

# ─── Pydantic schemas ───────────────────────────────────────────────────────


class SubmissionRequirement(BaseModel):
    category: str
    requirement: str
    mandatory: bool
    source_pdf: str
    source_page: int
    evidence: str

    @field_validator("source_page")
    @classmethod
    def page_must_be_positive(cls, v: int) -> int:
        if v < 1:
            raise ValueError("source_page must be >= 1")
        return v


class TradeAndScope(BaseModel):
    csi_division: str
    trade: str
    scope_description: str
    estimated_cost: Optional[float] = None
    vendor_name: Optional[str] = None
    quantity: Optional[str] = None
    source_pdf: str
    source_page: int
    evidence: str
    addendum_number: Optional[int] = None

    @field_validator("source_page")
    @classmethod
    def page_must_be_positive(cls, v: int) -> int:
        if v < 1:
            raise ValueError("source_page must be >= 1")
        return v


class PageExtraction(BaseModel):
    submission_requirements: list[SubmissionRequirement]
    trades_and_scope: list[TradeAndScope]
    flags: list[str] = []


# ─── Prompts ─────────────────────────────────────────────────────────────────

EXTRACTION_PROMPT = """\
You are a construction document analyst. You are looking at a single page from \
a construction PDF drawing set.

Your task: extract ONLY information that is clearly visible on THIS page.

Return a JSON object with exactly this structure (no markdown fences, no extra keys):
{{
  "submission_requirements": [
    {{
      "category": "<string: e.g. Permits, Submittals, Inspections, Insurance, Bonds, Certifications, etc.>",
      "requirement": "<string: concise description of the requirement>",
      "mandatory": <boolean: true if clearly mandatory, false if optional or unclear>,
      "source_pdf": "{pdf_name}",
      "source_page": {page_num},
      "evidence": "<string: short verbatim quote or paraphrase from the page proving this>"
    }}
  ],
  "trades_and_scope": [
    {{
      "csi_division": "<string: 2-digit CSI MasterFormat division if identifiable, else 'NA'>",
      "trade": "<string: trade name e.g. Earthwork, Concrete, Structural Steel, Plumbing, Electrical, HVAC, etc.>",
      "scope_description": "<string: brief description of scope visible on this page>",
      "estimated_cost": <number or null: dollar amount for this line item. Extract from bids, budgets, estimates, allowances, lump sums, unit prices x qty, contract amounts, SOV amounts, or any dollar figure tied to this trade. Strip $ and commas — e.g. $12,500 -> 12500. null if no dollar amount visible>,
      "vendor_name": "<string or null: name of the bidder, contractor, vendor, subcontractor, or supplier associated with this cost/trade — e.g. 'Eastern Concrete', 'L&R Mechanical', 'Budget'. null if no vendor/bidder name is visible>",
      "quantity": "<string or null: quantity with unit if visible, e.g. '2,400 SF', '150 LF', '12 EA'. null if none>",
      "source_pdf": "{pdf_name}",
      "source_page": {page_num},
      "evidence": "<string: short verbatim quote or visual description from the page proving this>"
    }}
  ]
}}

RULES:
- ONLY extract what is clearly visible. Do NOT hallucinate or infer beyond what is shown.
- If the page is a cover sheet, title block, or contains no extractable info, return empty arrays.
- For drawings: identify trades from drawing content (e.g., site plan implies Earthwork/Grading, \
  electrical symbols imply Electrical trade, plumbing fixtures imply Plumbing, etc.)
- For notes/specs: extract submission requirements and any trade/scope references.
- FINANCIAL DATA IS CRITICAL: This is a buyout/budgeting tool. Carefully scan every page for:
  * Dollar amounts ($), costs, prices, budgets, allowances, bids, estimates, contract values
  * Bid tabulations, cost breakdowns, SOV (Schedule of Values), line-item pricing
  * Unit prices (e.g., $4.50/SF) — multiply by quantity if both are shown
  * Lump-sum amounts, subtotals, per-trade totals
  * If a table or list shows trade names with dollar amounts, create one entry per trade with its cost
  Strip $ signs and commas — return a plain number (e.g. $12,500 -> 12500).
- VENDOR/BIDDER NAMES: If a company name, contractor, subcontractor, or supplier is associated \
  with a cost or trade, capture it in vendor_name. If the same trade has multiple vendors with \
  different bids, create SEPARATE entries for each vendor+amount pair.
- QUANTITIES: If quantities (SF, LF, EA, CY, SY, etc.) are visible, capture them in quantity.
- evidence must be a SHORT snippet (max ~80 chars) directly from the page.
- Return ONLY valid JSON. No markdown code fences. No commentary outside the JSON.
- SUBMISSION REQUIREMENTS — only extract items that meet at least ONE of:
  * Has a specific deadline, date, or schedule constraint
  * Involves money, retainage, bonds, or insurance amounts
  * Is a license, certification, permit, or regulatory approval requirement
  * Creates legal or schedule risk if missed (liquidated damages, inspections with teeth)
  * Is explicitly marked mandatory, "shall", "required", or "must"
  * Involves coordination with an outside authority (city, DOT, utility, engineer of record)
  Do NOT extract: general best practices, standard boilerplate phrases like "contractor
  shall verify all dimensions", informational notes with no legal consequence, or items
  that are repeated verbatim from another page already processed.
"""

RETRY_PROMPT = """\
The previous response was not valid JSON matching the required schema.

Return ONLY a valid JSON object with this exact structure:
{{
  "submission_requirements": [
    {{
      "category": "string",
      "requirement": "string",
      "mandatory": true/false,
      "source_pdf": "{pdf_name}",
      "source_page": {page_num},
      "evidence": "string"
    }}
  ],
  "trades_and_scope": [
    {{
      "csi_division": "string",
      "trade": "string",
      "scope_description": "string",
      "estimated_cost": null,
      "vendor_name": null,
      "quantity": null,
      "source_pdf": "{pdf_name}",
      "source_page": {page_num},
      "evidence": "string"
    }}
  ]
}}

If the page has no relevant content, return:
{{"submission_requirements": [], "trades_and_scope": []}}

ONLY output valid JSON. No markdown. No explanation.
"""


# ─── Helpers ─────────────────────────────────────────────────────────────────


def render_page_to_png(doc: fitz.Document, page_idx: int, dpi: int = DPI) -> bytes:
    """Render a single PDF page to PNG bytes at the given DPI.
    If the image exceeds MAX_IMAGE_BYTES, re-render at progressively lower DPI."""
    page = doc.load_page(page_idx)
    current_dpi = dpi
    while current_dpi >= 72:
        zoom = current_dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png_bytes = pix.tobytes("png")
        if len(png_bytes) <= MAX_IMAGE_BYTES:
            return png_bytes
        current_dpi = int(current_dpi * 0.75)
    return png_bytes


def extract_json_from_text(text: str) -> str:
    """Try to extract JSON from response text, stripping markdown fences if present."""
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
        text = text.strip()
    return text


def normalize_key(s: str) -> str:
    """Normalize a string for dedup comparison."""
    return re.sub(r"\s+", " ", s.strip().lower())


def dedup_requirements(items: list[SubmissionRequirement]) -> list[SubmissionRequirement]:
    """Remove duplicate requirements based on normalized category+requirement."""
    seen = set()
    result = []
    for item in items:
        key = (normalize_key(item.category), normalize_key(item.requirement))
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


def dedup_trades(items: list[TradeAndScope]) -> list[TradeAndScope]:
    """Remove duplicate trades based on normalized trade+scope_description."""
    seen = set()
    result = []
    for item in items:
        key = (normalize_key(item.trade), normalize_key(item.scope_description))
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


# ─── Addenda detection & conflict resolution ─────────────────────────────


_ADDENDUM_PATTERN = re.compile(
    r'(?:addendum|add)[- ._]*(?:no\.?|number|#)?[- ._]*(\d+)',
    re.IGNORECASE,
)


def detect_addendum_number(source_pdf: str, evidence: str, scope: str) -> Optional[int]:
    """Try to detect an addendum number from PDF name, evidence, or scope text."""
    for text in (source_pdf, evidence, scope):
        if not text:
            continue
        m = _ADDENDUM_PATTERN.search(text)
        if m:
            return int(m.group(1))
    return None


def tag_addenda(trades: list[TradeAndScope]) -> list[TradeAndScope]:
    """Tag each trade extraction with its addendum number (if detectable)."""
    for t in trades:
        if t.addendum_number is None:
            num = detect_addendum_number(t.source_pdf, t.evidence, t.scope_description)
            if num is not None:
                t.addendum_number = num
    return trades


def detect_addenda_conflicts(
    trades: list[TradeAndScope],
    emit: Callable[[str], None],
) -> list[dict]:
    """
    After all trades are extracted, detect specification conflicts between
    base documents and addenda. Highest addendum number wins.

    Returns a list of conflict finding dicts for the Flags tab.
    """
    # Group by trade + spec item (material/method keywords)
    # We look for the same trade with different scope details from different sources
    from collections import defaultdict

    trade_specs: dict[str, list[TradeAndScope]] = defaultdict(list)
    for t in trades:
        # Normalize trade name for grouping
        key = t.trade.strip().lower()
        trade_specs[key].append(t)

    findings = []
    for trade_key, items in trade_specs.items():
        # Only check trades with items from multiple sources or addenda
        addendum_items = [i for i in items if i.addendum_number is not None]
        base_items = [i for i in items if i.addendum_number is None]

        if not addendum_items or not base_items:
            continue

        # Check for cost conflicts (different dollar amounts)
        base_costs = [(i, i.estimated_cost) for i in base_items if i.estimated_cost]
        add_costs = [(i, i.estimated_cost) for i in addendum_items if i.estimated_cost]

        if base_costs and add_costs:
            # Highest addendum number wins
            winning = max(addendum_items, key=lambda x: x.addendum_number or 0)
            for base_item, base_cost in base_costs:
                if winning.estimated_cost and abs(winning.estimated_cost - base_cost) > 100:
                    finding = {
                        "severity": "HIGH",
                        "flag": "Addendum cost override",
                        "detail": (
                            f"CONFLICT: {trade_key.title()} cost "
                            f"${base_cost:,.0f} in {base_item.source_pdf} p.{base_item.source_page} "
                            f"vs ${winning.estimated_cost:,.0f} in "
                            f"Addendum {winning.addendum_number} "
                            f"({winning.source_pdf} p.{winning.source_page}). "
                            f"Addendum {winning.addendum_number} takes precedence. "
                            f"Using ${winning.estimated_cost:,.0f}. Please verify."
                        ),
                        "source": f"Addendum {winning.addendum_number}",
                    }
                    findings.append(finding)
                    emit(
                        f"[ADDENDUM] {trade_key.title()}: "
                        f"${base_cost:,.0f} -> ${winning.estimated_cost:,.0f} "
                        f"(Addendum {winning.addendum_number} overrides)"
                    )

        # Check for scope/spec text conflicts (different materials, methods)
        # Simple heuristic: if addendum scope mentions different materials
        for add_item in addendum_items:
            for base_item in base_items:
                # Check if scope descriptions mention conflicting specifics
                add_scope_lower = add_item.scope_description.lower()
                base_scope_lower = base_item.scope_description.lower()

                # Look for common spec keywords that differ
                _spec_markers = ['revised', 'replaced', 'changed', 'deleted',
                                 'in lieu of', 'substitute', 'modified']
                if any(marker in add_scope_lower for marker in _spec_markers):
                    finding = {
                        "severity": "MEDIUM",
                        "flag": "Addendum spec revision",
                        "detail": (
                            f"CONFLICT: {trade_key.title()} \u2014 "
                            f"Base: \"{base_item.scope_description[:80]}\" "
                            f"({base_item.source_pdf} p.{base_item.source_page}) "
                            f"vs Addendum {add_item.addendum_number}: "
                            f"\"{add_item.scope_description[:80]}\" "
                            f"({add_item.source_pdf} p.{add_item.source_page}). "
                            f"Addendum {add_item.addendum_number} takes precedence."
                        ),
                        "source": f"Addendum {add_item.addendum_number}",
                    }
                    findings.append(finding)
                    emit(
                        f"[ADDENDUM] {trade_key.title()}: spec revision in "
                        f"Addendum {add_item.addendum_number}"
                    )
                    break  # One finding per addendum item is enough

    if findings:
        emit(f"Found {len(findings)} addenda conflict(s)")
    return findings


# ─── Supported file types ──────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    # PDFs (image-based extraction)
    ".pdf",
    # Images (direct to Vision API)
    ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp",
    # Documents (text extraction)
    ".docx", ".doc",
    # Spreadsheets (text extraction)
    ".xlsx", ".xls", ".csv",
    # Email (text extraction)
    ".eml",
    # Plain text
    ".txt", ".rtf",
}

# Max characters per text chunk sent to Claude
_TEXT_CHUNK_SIZE = 8000


def _extract_text_from_docx(file_path: Path) -> str:
    """Extract all text from a DOCX file."""
    if DocxDocument is None:
        raise ImportError("python-docx is required for DOCX files. Install with: pip install python-docx")
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_text_from_eml(file_path: Path) -> str:
    """Extract text content from an EML email file."""
    with open(file_path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    parts = []
    # Headers
    for header in ["From", "To", "Subject", "Date", "CC"]:
        val = msg.get(header)
        if val:
            parts.append(f"{header}: {val}")

    parts.append("")  # blank line after headers

    # Body
    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        content = body.get_content()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        parts.append(content)

    # List attachments
    attachments = []
    for part in msg.walk():
        fn = part.get_filename()
        if fn:
            attachments.append(fn)
    if attachments:
        parts.append("\n--- Attachments ---")
        for att in attachments:
            parts.append(f"  - {att}")

    return "\n".join(parts)


def _extract_text_from_spreadsheet(file_path: Path) -> str:
    """Extract text content from XLSX, XLS, or CSV files."""
    ext = file_path.suffix.lower()
    parts = []

    if ext == ".csv":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row_idx, row in enumerate(reader):
                if any(cell.strip() for cell in row):
                    parts.append(" | ".join(cell.strip() for cell in row))
    else:
        # XLSX / XLS via pandas
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
                parts.append(f"--- Sheet: {sheet_name} ---")
                for _, row in df.iterrows():
                    cells = [str(v).strip() for v in row if pd.notna(v) and str(v).strip()]
                    if cells:
                        parts.append(" | ".join(cells))
                parts.append("")
        except Exception as e:
            parts.append(f"[Error reading spreadsheet: {e}]")

    return "\n".join(parts)


def _extract_text_from_txt(file_path: Path) -> str:
    """Read plain text or RTF file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _chunk_text(text: str, chunk_size: int = _TEXT_CHUNK_SIZE) -> list[str]:
    """Split text into chunks, breaking at line boundaries."""
    lines = text.split("\n")
    chunks = []
    current = []
    current_len = 0

    for line in lines:
        line_len = len(line) + 1
        if current_len + line_len > chunk_size and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(line)
        current_len += line_len

    if current:
        chunks.append("\n".join(current))

    return chunks if chunks else ["(empty document)"]


def _get_image_media_type(file_path: Path) -> str:
    """Get MIME type for image files."""
    ext = file_path.suffix.lower()
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".tiff": "image/tiff",
        ".tif": "image/tiff",
        ".bmp": "image/bmp",
        ".webp": "image/webp",
    }.get(ext, "image/png")


# ─── Text-based page processing ───────────────────────────────────────────


def process_text_page(
    client: anthropic.Anthropic,
    text_content: str,
    file_name: str,
    page_num: int,
    stats: dict,
) -> Optional[PageExtraction]:
    """Send text content (from DOCX/EML/XLSX/etc.) to Claude for extraction."""

    prompt_text = EXTRACTION_PROMPT.format(pdf_name=file_name, page_num=page_num)

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"--- DOCUMENT CONTENT ({file_name}, section {page_num}) ---\n\n"
                            f"{text_content}\n\n"
                            f"--- END DOCUMENT CONTENT ---\n\n"
                            f"{prompt_text}",
                },
            ],
        }
    ]

    SIMPLIFIED_PROMPT = """
This is a construction document section.
Extract ANY trades or work items visible.
Return valid JSON only:
{
  "trades_and_scope": [
    {"trade": "trade name", "scope": ["scope item"]}
  ],
  "submission_requirements": [],
  "flags": []
}
If truly nothing is extractable, return:
{"trades_and_scope": [], "submission_requirements": [], "flags": []}
"""

    for attempt in range(4):
        try:
            # On attempts 3 and 4, use simplified prompt
            if attempt >= 2:
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": f"--- DOCUMENT CONTENT ({file_name}, section {page_num}) ---\n\n"
                                        f"{text_content}\n\n"
                                        f"--- END DOCUMENT CONTENT ---\n\n"
                                        f"{SIMPLIFIED_PROMPT}",
                            },
                        ],
                    }
                ]

            response = client.messages.create(
                model=MODEL,
                max_tokens=4096,
                temperature=TEMPERATURE,
                messages=messages,
            )

            usage = response.usage
            stats["input_tokens"] += usage.input_tokens
            stats["output_tokens"] += usage.output_tokens
            stats["api_calls"] += 1

            raw_text = response.content[0].text
            json_text = extract_json_from_text(raw_text)

            data = json.loads(json_text)
            extraction = PageExtraction(**data)
            return extraction

        except (json.JSONDecodeError, ValidationError):
            if attempt == 0:
                retry_text = RETRY_PROMPT.format(pdf_name=file_name, page_num=page_num)
                messages.append({"role": "assistant", "content": raw_text})
                messages.append({"role": "user", "content": retry_text})
            elif attempt == 3:
                stats["failed_pages"].append(f"{file_name} section {page_num}")

        except anthropic.APIError:
            if attempt < 3:
                time.sleep(5)
            else:
                stats["failed_pages"].append(f"{file_name} section {page_num}")

    # Return empty but valid result — never lose a page silently
    return PageExtraction(
        trades_and_scope=[],
        submission_requirements=[],
        flags=[f"WARNING: Section {page_num} could not be parsed after 4 attempts — manual review required"]
    )


# ─── Trade consolidation ────────────────────────────────────────────────────

# Master trade list matching the buyout sheet format.
# Each entry: (sort_order, division, trade_name, keyword_patterns)
# keyword_patterns match against BOTH the raw trade name and the CSI division.
# First match wins, so more specific patterns go first.
_BUYOUT_TRADES: list[tuple[int, str, str, list[str]]] = [
    # Div 2 — Sitework
    (1,  "2",  "Sitework",               ["sitework", "earthwork", "grading", "excavat", "erosion",
                                           "sediment", "site drainage", "demolit", "abatement",
                                           "existing condition", "drainage", "storm drain",
                                           "catch basin", "stormwater", "utilities", "water util",
                                           "sewer util", "site improv", "exterior improv",
                                           "retaining wall", "erosion control"]),
    (2,  "2",  "Paving & Striping",      ["paving", "striping", "asphalt", "parking"]),
    (3,  "2",  "Landscaping",            ["landscape", "landscaping", "planting", "irrigation",
                                           "bioretention", "plant list"]),
    (4,  "2",  "Fencing",                ["fenc"]),
    # Div 3 — Concrete
    (5,  "3",  "Site Concrete",          ["site concrete", "sidewalk", "curb", "ramp",
                                           "concrete flatwork"]),
    (6,  "3",  "Building Concrete",      ["concrete", "foundation", "footing", "slab"]),
    (29, "4",  "Masonry",                ["masonry", "cmu", "cmu block", "cast stone",
                                           "brick veneer", "brick wall",
                                           "stonework", "stone veneer",
                                           "block wall", "block veneer"]),
    # Div 5 — Metals
    (7,  "5",  "Structural Steel",       ["structural", "steel", "metal connect", "simpson",
                                           "strap", "clip", "anchor bolt"]),
    # Div 6 — Wood & Plastics
    (8,  "6",  "Rough Carpentry",        ["rough carp", "wood fram", "framing", "stud", "joist",
                                           "rafter", "truss", "sheathing", "blocking",
                                           "roof framing", "attic framing", "wood/plastic",
                                           "composites", "pvc", "architectural",
                                           "equipment plan"]),
    (10, "8",  "Windows",                ["window", "glazing", "marvin"]),
    (9,  "8",  "Doors/Hdwr/Finish Carp", ["door", "hardware", "finish carp", "trim", "casework",
                                           "millwork", "woodwork", "opening",
                                           "hdwr", "lockset", "lever handle"]),
    # Div 7 — Thermal & Moisture
    (11, "7",  "Insulation",             ["insulation", "insul", "spray foam", "batt",
                                           "thermal", "moisture protect", "vapor barrier",
                                           "air barrier", "air infiltrat", "moisture bar",
                                           "housewrap", "tyvek", "hydrogap"]),
    (14, "5",  "Gutters",               ["gutter", "downspout"]),
    (15, "7",  "Flashing & Waterproofing", ["waterproof", "dampproof", "damp-proof",
                                           "flash", "sheet metal", "drip edge",
                                           "sealant", "caulk", "backer rod", "termite"]),
    (12, "7",  "Roofing",               ["roofing", "shingle", "roof assembly", "roof membran",
                                           "ice.*water", "underlayment", "roof felt"]),
    (13, "7",  "Siding",                ["siding", "exterior clad", "vinyl sid",
                                           "exterior enclos", "cedar impression"]),
    # Div 9 — Finishes
    (16, "9",  "Drywall",               ["drywall", "gypsum", "gwb", "resilient channel"]),
    (17, "9",  "Flooring",              ["floor", "lvt", "vinyl tile", "carpet", "tile",
                                           "ceramic", "non-slip"]),
    (18, "9",  "Painting",              ["paint", "coating", "primer", "finish schedule",
                                           "room finish"]),
    (19, "9",  "Countertops & Finishes", ["countertop", "hi-macs", "finish"]),
    # Div 10-13 — Specialties & Equipment
    (20, "10", "Specialties",            ["specialt", "accessori", "grab bar", "mirror",
                                           "towel bar", "shower control"]),
    # Div 14 — Conveying
    (22, "14", "Conveying Equipment",    ["elevator", "conveying", "lift", "ceiling lift"]),
    # Div 15 (22-23) — Mechanical (plumbing BEFORE cabinets so "Plumbing/Kitchen" matches here)
    (23, "15", "Plumbing",              ["plumb", "piping", "fixture", "toilet", "lavatory",
                                           "shower", "bathtub", "water heater", "vanity",
                                           "penguin", "overflow"]),
    (24, "15", "HVAC",                  ["hvac", "mechanical", "mini-split", "mitsubishi",
                                           "air condition", "heating", "duct", "ventilat",
                                           "continuous vent"]),
    (28, "15", "Fire Sprinkler",        ["sprinkler", "fire protect", "fire suppress",
                                           "nfpa 13", "nfpa13", "wet pipe", "dry pipe",
                                           "fire suppression", "standpipe"]),
    # Cabinets after Plumbing so "Plumbing/Kitchen" doesn't match "kitchen" first
    (21, "13", "Cabinets",              ["cabinet", "kitchen", "furnish", "appliance",
                                           "express kitchen"]),
    # Div 16 (26-28) — Electrical
    (25, "16", "Electrical",            ["electric", "lighting", "outlet", "panel", "amp",
                                           "low voltage", "control panel", "site light",
                                           "luminaire", "wiring", "transformer"]),
    (26, "16", "Smart Home / Security",  ["communicat", "smart", "speaker", "network", "data",
                                           "safety", "security", "alarm", "arms", "fire protect",
                                           "camera", "doorbell", "sensor", "blind"]),
    # Div 1 — General (catch-all, placed last)
    (27, "1",  "General Requirements",   ["general", "project info", "coordination"]),
]


def _classify_trade(raw_trade: str, raw_csi: str) -> tuple[int, str, str]:
    """
    Classify a raw trade extraction into a buyout-style trade category.
    Returns (sort_order, division, canonical_trade_name).
    """
    lower = raw_trade.lower()

    for sort_order, div, canon, patterns in _BUYOUT_TRADES:
        for pat in patterns:
            if pat in lower:
                return (sort_order, div, canon)

    # Fallback: try to use the CSI division number to find a default bucket
    raw_num = str(raw_csi).strip().split(".")[0]
    try:
        num = int(float(raw_num))
    except (ValueError, TypeError):
        num = 0

    # Map extended CSI (22-33) to traditional divisions and find a default
    div_defaults = {
        1: (27, "1", "General Requirements"),
        2: (1, "2", "Sitework"),
        3: (6, "3", "Building Concrete"),
        4: (6, "3", "Building Concrete"),
        5: (7, "5", "Structural Steel"),
        6: (8, "6", "Rough Carpentry"),
        7: (11, "7", "Insulation"),
        8: (9, "8", "Doors/Hdwr/Finish Carp"),
        9: (16, "9", "Drywall"),
        10: (20, "10", "Specialties"),
        12: (21, "13", "Cabinets"),
        13: (21, "13", "Cabinets"),
        14: (22, "14", "Conveying Equipment"),
        21: (28, "15", "Fire Sprinkler"),
        22: (23, "15", "Plumbing"),
        23: (24, "15", "HVAC"),
        26: (25, "16", "Electrical"),
        27: (26, "16", "Smart Home / Security"),
        28: (26, "16", "Smart Home / Security"),
        31: (1, "2", "Sitework"),
        32: (2, "2", "Paving & Striping"),
        33: (1, "2", "Sitework"),
    }

    if num in div_defaults:
        return div_defaults[num]

    return (27, "1", "General Requirements")


def consolidate_trades(trades: list[TradeAndScope]) -> list[dict]:
    """
    Group raw per-page trade extractions into consolidated buyout-style rows:
    one row per canonical trade, with scope items listed in page order.

    Returns a list of dicts sorted by buyout order (matching the user's
    buyout sheet: Sitework, Paving, Landscaping, Concrete, ... Electrical).
    Each dict includes 'bids' — a list of (vendor_name, amount) pairs for
    populating the bid columns in the Excel output.
    """

    groups: dict[str, dict] = {}  # keyed by canonical trade name

    for t in trades:
        sort_order, div, canon = _classify_trade(t.trade, t.csi_division)
        key = canon

        if key not in groups:
            groups[key] = {
                "sort_order": sort_order,
                "csi_division": div,
                "trade": canon,
                "scope_items": [],       # (page_order_key, description, cost, qty)
                "sources": [],           # "PDF p.N"
                "seen_scope": set(),     # for dedup
                "total_cost": 0.0,       # sum of all extracted costs
                "has_cost": False,       # whether any cost was found
                "quantities": [],        # collected quantity strings
                "bids": {},              # vendor_name -> best (lowest) amount
                "unnamed_costs": [],     # costs with no vendor name
            }

        g = groups[key]

        # Dedup scope descriptions
        norm_scope = normalize_key(t.scope_description)
        if norm_scope not in g["seen_scope"]:
            g["seen_scope"].add(norm_scope)
            order_key = (t.source_pdf, t.source_page)
            g["scope_items"].append((order_key, t.scope_description,
                                     t.estimated_cost, t.quantity))

        # Collect bids: vendor + amount pairs
        if t.estimated_cost is not None and t.estimated_cost > 0:
            g["total_cost"] += t.estimated_cost
            g["has_cost"] = True

            if t.vendor_name and t.vendor_name.strip():
                vname = t.vendor_name.strip()
                # Keep the largest amount per vendor (they may bid on multiple scope items)
                if vname in g["bids"]:
                    g["bids"][vname] += t.estimated_cost
                else:
                    g["bids"][vname] = t.estimated_cost
            else:
                g["unnamed_costs"].append(t.estimated_cost)

        # Collect quantities
        if t.quantity:
            if t.quantity not in g["quantities"]:
                g["quantities"].append(t.quantity)

        source_ref = f"{t.source_pdf} p.{t.source_page}"
        if source_ref not in g["sources"]:
            g["sources"].append(source_ref)

    # Build consolidated rows sorted by buyout order
    result = []
    for key in sorted(groups.keys(), key=lambda k: groups[k]["sort_order"]):
        g = groups[key]
        # Sort scope items by page order (PDF name, then page number)
        g["scope_items"].sort(key=lambda x: x[0])
        # Build numbered scope text (include cost/qty inline when available)
        scope_lines = []
        for i, (_, desc, cost, qty) in enumerate(g["scope_items"], 1):
            parts = [f"{i}. {desc}"]
            if qty:
                parts.append(f"[{qty}]")
            if cost and cost > 0:
                parts.append(f"(${cost:,.0f})")
            scope_lines.append(" ".join(parts))
        scope_text = "\n".join(scope_lines)

        # Build bid list: sorted by amount (lowest first)
        bid_list = sorted(g["bids"].items(), key=lambda x: x[1])
        # If we have unnamed costs but no vendor bids, aggregate them as "Budget"
        if not bid_list and g["unnamed_costs"]:
            total_unnamed = sum(g["unnamed_costs"])
            bid_list = [("Budget", total_unnamed)]

        # Determine budget/SOV: use total_cost if any costs were found
        budget = g["total_cost"] if g["has_cost"] else None

        result.append({
            "csi_division": g["csi_division"],
            "trade": g["trade"],
            "scope": scope_text,
            "source_pages": "; ".join(g["sources"]),
            "budget": budget,
            "quantities": "; ".join(g["quantities"]) if g["quantities"] else None,
            "bids": bid_list,  # list of (vendor_name, amount) sorted low->high
        })

    return result


# ─── SOV mapping from reference buyout ──────────────────────────────────────


def parse_sov_from_buyout(file_path: Path) -> dict:
    """
    Parse SOV (Schedule of Values) data from a reference buyout spreadsheet.

    Searches for an 'SOV' column header, then reads trade name -> SOV value pairs.
    Also extracts summary values (GC OH&P, Bond Premium, Permit Fees).

    Returns a dict with:
        'trade_sov': dict[str, float] - reference trade name -> SOV value
        'ohp': float or None - GC OH&P value
        'bond': float or None - Bond Premium value
        'permit': float or None - Permit Fees value
    """
    df_dict = pd.read_excel(file_path, sheet_name=None, header=None)

    result = {"trade_sov": {}, "ohp": None, "bond": None, "permit": None}

    for sheet_name, df in df_dict.items():
        # Search for ALL "SOV" column headers, use the rightmost (most recent)
        sov_positions = []
        for col_idx in range(len(df.columns)):
            for row_idx in range(min(15, len(df))):
                val = df.iloc[row_idx, col_idx]
                if pd.notna(val) and isinstance(val, str) and "sov" in val.lower():
                    sov_positions.append((row_idx, col_idx))
                    break

        if not sov_positions:
            continue

        sov_header_row, sov_col = max(sov_positions, key=lambda x: x[1])

        # Find the trade name column: look for columns with trade-like names
        trade_col = None
        trade_keywords = [
            "site", "pav", "concrete", "carp", "insul", "roof", "sid",
            "door", "window", "drywall", "floor", "paint", "plumb",
            "hvac", "electric", "cabinet", "gutter", "landscape", "fenc",
            "steel", "masonry", "framing",
        ]
        best_count = 0
        for col_idx in range(min(10, len(df.columns))):
            trade_count = 0
            for row_idx in range(sov_header_row + 1, min(len(df), sov_header_row + 30)):
                val = df.iloc[row_idx, col_idx]
                if pd.notna(val) and isinstance(val, str) and len(val.strip()) > 2:
                    lower = val.lower().strip()
                    if any(kw in lower for kw in trade_keywords):
                        trade_count += 1
            if trade_count > best_count:
                best_count = trade_count
                trade_col = col_idx

        if trade_col is None:
            trade_col = min(1, len(df.columns) - 1)

        # Extract trade -> SOV pairs
        for row_idx in range(sov_header_row + 1, len(df)):
            trade_val = df.iloc[row_idx, trade_col]
            sov_val = df.iloc[row_idx, sov_col]

            if pd.isna(trade_val) or pd.isna(sov_val):
                continue

            trade_name = str(trade_val).strip()
            if not trade_name:
                continue

            try:
                sov_num = float(sov_val)
            except (ValueError, TypeError):
                continue

            if sov_num == 0:
                continue

            # Detect summary rows
            lower_trade = trade_name.lower()
            if "oh&p" in lower_trade or "overhead" in lower_trade:
                result["ohp"] = sov_num
            elif "bond" in lower_trade:
                result["bond"] = sov_num
            elif "permit" in lower_trade:
                result["permit"] = sov_num
            elif "total" in lower_trade or "contract" in lower_trade:
                pass  # Skip total/subtotal rows
            else:
                result["trade_sov"][trade_name] = sov_num

        # Only process the first sheet with SOV data
        if result["trade_sov"]:
            break

    return result


def _match_sov_to_trades(
    consolidated: list[dict], sov_mapping: dict[str, float]
) -> dict[str, float]:
    """
    Match reference buyout trade names to BuildBrain canonical trade names.

    Returns a dict of canonical_trade_name -> SOV value.
    Uses three-pass matching: exact, containment, keyword classification.
    """
    matched = {}
    used_refs = set()

    # Pass 1: Exact match (case-insensitive)
    for row in consolidated:
        canon = row["trade"]
        canon_lower = canon.lower().strip()
        for ref_trade, sov_val in sov_mapping.items():
            if ref_trade in used_refs:
                continue
            if ref_trade.lower().strip() == canon_lower:
                matched[canon] = sov_val
                used_refs.add(ref_trade)
                break

    # Pass 2: Containment match
    for row in consolidated:
        canon = row["trade"]
        if canon in matched:
            continue
        canon_lower = canon.lower().strip()
        for ref_trade, sov_val in sov_mapping.items():
            if ref_trade in used_refs:
                continue
            ref_lower = ref_trade.lower().strip()
            if canon_lower in ref_lower or ref_lower in canon_lower:
                matched[canon] = sov_val
                used_refs.add(ref_trade)
                break

    # Pass 3: Classify reference trade using same keyword patterns
    for row in consolidated:
        canon = row["trade"]
        if canon in matched:
            continue
        for ref_trade, sov_val in sov_mapping.items():
            if ref_trade in used_refs:
                continue
            _, _, ref_canonical = _classify_trade(ref_trade, "")
            if ref_canonical == canon:
                matched[canon] = sov_val
                used_refs.add(ref_trade)
                break

    return matched


# ─── Pricing engine ──────────────────────────────────────────────────────────


def calculate_trade_estimate(
    trade_name: str,
    work_items: list[dict],
    wage_regime: str = "CT_DOL_RESIDENTIAL",
    sub_quote: float = None,
    sub_quote_source: str = None
) -> dict:
    """
    Calculates a complete trade estimate with full calculation trace.

    work_items is a list of dicts, each with:
        - "work_item": key from PRODUCTIVITY_RATES
        - "quantity": numeric quantity
        - "material_key": key from MATERIAL_DB (optional)
        - "material_quantity": quantity of material (optional, defaults to same as work quantity)
        - "description": human-readable description

    If sub_quote is provided, it OVERRIDES the calculated baseline.
    The calculated baseline is still shown for comparison.

    Returns a dict with full trace suitable for Excel SOV output.
    """
    line_items = []
    total_material = 0.0
    total_labor = 0.0

    for item in work_items:
        work_item_key = item.get("work_item")
        quantity = item.get("quantity", 0)
        material_key = item.get("material_key")
        mat_qty = item.get("material_quantity", quantity)
        description = item.get("description", work_item_key)

        # Labor calculation
        labor_result = None
        if work_item_key and quantity > 0:
            labor_result = calculate_labor(work_item_key, quantity, wage_regime)

        # Material calculation
        material_result = None
        if material_key and mat_qty > 0:
            mat_data = get_material_price(material_key)
            if mat_data:
                mat_cost = mat_qty * mat_data["price"]
                material_result = {
                    "key": material_key,
                    "quantity": mat_qty,
                    "unit": mat_data["unit"],
                    "unit_price": mat_data["price"],
                    "cost": mat_cost,
                    "note": mat_data["note"],
                    "trace": f"{mat_qty} {mat_data['unit']} \u00d7 ${mat_data['price']:.2f}/{mat_data['unit']} = ${mat_cost:,.0f}"
                }
                total_material += mat_cost

        labor_cost = labor_result["labor_cost"] if labor_result else 0
        total_labor += labor_cost

        line_items.append({
            "description": description,
            "labor": labor_result,
            "material": material_result,
            "line_total": labor_cost + (material_result["cost"] if material_result else 0)
        })

    calculated_total = total_material + total_labor

    # Sub quote override logic
    if sub_quote is not None:
        variance = sub_quote - calculated_total
        variance_pct = (variance / calculated_total * 100) if calculated_total > 0 else 0
        use_amount = sub_quote
        override_flag = True
        override_note = f"SUB QUOTE USED: ${sub_quote:,.0f} from {sub_quote_source or 'unnamed source'}. Baseline was ${calculated_total:,.0f} (variance: {variance_pct:+.1f}%)"
        if abs(variance_pct) > 20:
            override_flag = "REVIEW"  # flag for human review if >20% off
    else:
        use_amount = calculated_total
        override_flag = False
        override_note = "Calculated baseline \u2014 no sub quote provided"

    return {
        "trade": trade_name,
        "wage_regime": wage_regime,
        "wage_regime_label": CT_WAGE_RATES.get(wage_regime, {}).get("label", wage_regime),
        "line_items": line_items,
        "total_material": round(total_material, 2),
        "total_labor": round(total_labor, 2),
        "calculated_total": round(calculated_total, 2),
        "sub_quote": sub_quote,
        "sub_quote_source": sub_quote_source,
        "use_amount": round(use_amount, 2),
        "sub_quote_override": override_flag,
        "override_note": override_note,
        "confidence": "HIGH" if sub_quote else "MEDIUM",
        "accuracy_note": "Sub quote from named source" if sub_quote else "Calculated baseline \u00b110-15%"
    }


def detect_wage_regime(project_data: dict) -> dict:
    """
    Determines the applicable wage regime based on project characteristics.
    Returns {"regime": str, "confidence": str, "flags": list[str]}

    project_data should contain any of: owner, funding_source, project_type,
    location, stories, notes
    """
    flags = []
    regime = "CT_DOL_RESIDENTIAL"  # safe default
    confidence = "MEDIUM"

    text_to_check = " ".join([
        str(project_data.get("owner", "")),
        str(project_data.get("funding_source", "")),
        str(project_data.get("notes", "")),
        str(project_data.get("project_type", "")),
    ]).lower()

    # Federal funding triggers Davis-Bacon
    federal_keywords = ["hud", "home program", "cdbg", "lihtc", "fha",
                        "usda", "federal", "davis-bacon", "davis bacon"]
    for kw in federal_keywords:
        if kw in text_to_check:
            regime = "DAVIS_BACON_BUILDING_NEW_HAVEN"
            confidence = "HIGH"
            flags.append(f"Federal funding keyword detected: '{kw}' \u2014 Davis-Bacon applies")
            break

    # Non-profit / housing authority suggests prevailing wage likely
    nonprofit_keywords = ["arc", "housing authority", "affordable", "nonprofit",
                          "non-profit", "501c", "community development"]
    for kw in nonprofit_keywords:
        if kw in text_to_check:
            flags.append(f"Non-profit/affordable housing indicator: '{kw}' \u2014 confirm funding source for Davis-Bacon applicability")
            if confidence != "HIGH":
                confidence = "LOW"
            break

    # CT DOT / state roadway work
    dot_keywords = ["ct dot", "dot form 816", "state roadway", "public right of way", "state highway"]
    for kw in dot_keywords:
        if kw in text_to_check:
            flags.append(f"CT DOT / state roadway work detected \u2014 verify prevailing wage for ROW work")
            break

    # Residential building type (confirms CT DOL residential is correct)
    if any(k in text_to_check for k in ["single family", "residential", "apartment", "housing"]):
        if regime == "CT_DOL_RESIDENTIAL":
            confidence = "HIGH"

    flags.append(f"Annual adjustment: CT DOL rates update July 1 each year \u2014 verify before bid")

    return {
        "regime": regime,
        "regime_label": CT_WAGE_RATES.get(regime, {}).get("label", regime),
        "confidence": confidence,
        "flags": flags
    }


def detect_project_type(extraction: dict) -> dict:
    """
    Detects whether this is a standard residential project or a
    specialty/ADA accessible housing project.

    Returns:
    {
        "project_type": "STANDARD_RESIDENTIAL" | "ADA_ACCESSIBLE" | "COMMERCIAL",
        "labor_multiplier": float,  # 1.0 for standard, 1.35 for ADA
        "confidence": "HIGH" | "MEDIUM" | "LOW",
        "trigger_items": list[str],  # what triggered the classification
        "flag_message": str
    }
    """

    ADA_KEYWORDS = [
        "ceiling lift", "grab bar", "roll-in shower", "roll under",
        "arms monitoring", "arms system", "panic button", "nurse call",
        "body dryer", "smart bed", "hospital grade", "medical grade",
        "touchless faucet", "auto dispenser", "accessibility",
        "ada compliance", "ada compliant", "wheelchair", "accessible",
        "supported living", "assisted living", "group home",
        "hydrogiene", "bidette", "shower chair", "transfer bench",
        "hearing loop", "visual alarm", "tactile"
    ]

    SPECIALTY_EQUIPMENT_KEYWORDS = [
        "arms", "ceiling track", "lift system", "smart home integration",
        "remote monitoring", "sensor", "auto", "touchless", "voice control"
    ]

    # Pull all text from extraction to search
    all_text = " ".join([
        str(extraction.get("project_summary", "")),
        str(extraction.get("scope", "")),
        str(extraction.get("equipment_list", "")),
        str(extraction.get("special_requirements", "")),
        str(extraction.get("notes", "")),
        str(extraction.get("owner", "")),
        str(extraction.get("raw_text", "")),
    ]).lower()

    trigger_items = []

    for kw in ADA_KEYWORDS:
        if kw in all_text:
            trigger_items.append(kw)

    specialty_count = sum(1 for kw in SPECIALTY_EQUIPMENT_KEYWORDS if kw in all_text)

    # Count specialty equipment items if equipment list exists
    equipment_list = extraction.get("equipment_list", [])
    equipment_count = len(equipment_list) if isinstance(equipment_list, list) else 0

    # Classification logic
    if len(trigger_items) >= 3 or equipment_count >= 15 or specialty_count >= 3:
        project_type = "ADA_ACCESSIBLE"
        labor_multiplier = 1.35
        confidence = "HIGH" if len(trigger_items) >= 5 else "MEDIUM"
        flag_message = (
            f"ADA/ACCESSIBLE HOUSING DETECTED \u2014 {len(trigger_items)} specialty "
            f"indicators found: {', '.join(trigger_items[:5])}{'...' if len(trigger_items) > 5 else ''}. "
            f"Labor costs adjusted by 1.35x to reflect specialty installation "
            f"requirements, extended coordination time, and ADA-grade workmanship. "
            f"Material costs unchanged. Estimator should verify multiplier."
        )
    elif len(trigger_items) >= 1 or equipment_count >= 8:
        project_type = "ADA_ACCESSIBLE"
        labor_multiplier = 1.20
        confidence = "LOW"
        flag_message = (
            f"POSSIBLE SPECIALTY SCOPE \u2014 {len(trigger_items)} ADA/specialty "
            f"indicators found. Labor adjusted by 1.20x as precaution. "
            f"Estimator must review and confirm project type."
        )
    else:
        project_type = "STANDARD_RESIDENTIAL"
        labor_multiplier = 1.0
        confidence = "HIGH"
        flag_message = "Standard residential scope detected. No ADA labor premium applied."

    return {
        "project_type": project_type,
        "labor_multiplier": labor_multiplier,
        "confidence": confidence,
        "trigger_items": trigger_items,
        "flag_message": flag_message
    }


def get_baseline_quantities(trade_name: str, extraction: dict,
                            project_quantities: dict = None) -> list[dict]:
    """
    Maps a trade name to work items with quantities.
    Quantities are derived from project dimensions extracted from drawings
    via extract_project_quantities(), with conservative fallbacks.

    Returns a list of work_item dicts for calculate_trade_estimate().
    Returns empty list if trade cannot be auto-priced.
    """

    q = project_quantities or {}

    # Core dimensions — extracted or estimated
    unit_count    = q.get('unit_count') or 1
    floor_count   = q.get('floor_count') or 1

    # Type-aware SF defaults (replaces hardcoded 699)
    _default_sf = {'single_family': 1800, 'multi_family': 10000,
                   'commercial': 8000, 'mixed_use': 12000}
    project_type = q.get('project_type', 'unknown')
    _fallback_total_sf = _default_sf.get(project_type, 2000)

    footprint_sf  = q.get('footprint_sf') or \
                    (q.get('total_building_sf') or _fallback_total_sf) / max(floor_count, 1)
    total_sf      = q.get('total_building_sf') or \
                    footprint_sf * floor_count
    perimeter_lf  = q.get('perimeter_lf') or \
                    (footprint_sf ** 0.5) * 4  # square approximation

    # Project type detection for system routing
    project_type = q.get('project_type', 'unknown')
    is_multifamily = project_type in ('multi_family', 'mixed_use', 'commercial', 'institutional')
    is_commercial = project_type in ('commercial', 'mixed_use', 'institutional')
    is_residential_units = (
        is_multifamily and
        unit_count is not None and
        unit_count > 0 and
        project_type in ('multi_family', 'mixed_use')
    )
    # institutional/community_center: no per-unit formulas

    # Unit type for per-unit scaling
    unit_type = q.get('unit_type', 'standard')
    if unit_type == 'standard' and unit_count > 15:
        unit_type = 'sro'  # likely SRO/studio if high unit count in small building

    # Scale envelope by total building SF not just footprint
    floor_height  = q.get('floor_to_floor_height_ft') or 9
    ext_wall_sf   = perimeter_lf * floor_count * floor_height
    roof_sq       = (footprint_sf * 1.15) / 100     # 15% pitch factor

    # Scale MEP by unit count — Fix 8: window count by project type
    if is_multifamily:
        # SRO/studio: 2-3 windows, standard 1BR: 4-5, 2BR+: 6-8
        if unit_type == 'sro' or unit_count > 15:
            _win_per_unit = 3
        elif unit_count > 6:
            _win_per_unit = 5
        else:
            _win_per_unit = 8
        window_count = q.get('window_count') or (unit_count * _win_per_unit)
    else:
        window_count = q.get('window_count') or unit_count * 8
    ext_door_count    = q.get('ext_door_count') or max(2, unit_count // 10)
    int_door_count    = q.get('int_door_count') or unit_count * 4
    plumbing_fixtures = q.get('plumbing_fixtures') or unit_count * 5
    hvac_zones        = q.get('hvac_zones') or unit_count
    electrical_circuits = (q.get('circuits_per_unit') or 18) * unit_count
    electrical_panels = q.get('electrical_panels') or \
                        max(1, unit_count // 42)

    # Scale finishes by total SF
    cabinet_lf    = unit_count * 16       # 16 LF per unit
    tile_sf       = unit_count * 60       # 60 SF tile per unit
    lvt_sf        = total_sf - tile_sf    # rest is LVT
    foundation_lf = perimeter_lf

    # Interior walls scale with total SF and unit count — Fix 4 & 5
    if is_multifamily:
        # Multi-unit: unit separation walls (fire-rated), corridors, shafts
        avg_unit_perimeter = (total_sf / max(unit_count, 1)) ** 0.5 * 4
        unit_separation_sf = unit_count * avg_unit_perimeter * floor_height
        corridor_wall_sf = total_sf * 0.15 * floor_height  # 15% of floor is corridor
        shaft_wall_sf = floor_count * 4 * 30 * floor_height  # elevator + stair shafts
        int_wall_sf = (total_sf * 0.6) + unit_separation_sf + corridor_wall_sf + shaft_wall_sf
    else:
        int_wall_sf = total_sf * 0.9        # 0.9x total SF (residential)

    if is_multifamily:
        # Multi-unit: both sides of walls + ceiling + fire-rated corridor ceilings
        common_area_ceiling_sf = total_sf * 0.25  # corridors, lobbies, common areas
        gwb_sf = (int_wall_sf * 2) + ext_wall_sf + total_sf + (common_area_ceiling_sf * 1.5)
    else:
        gwb_sf = (int_wall_sf * 2) + ext_wall_sf + total_sf

    TRADE_MAP = {
        "Building Concrete": [
            {"work_item": "slab_on_grade_4in",    "quantity": footprint_sf,
             "material_key": "concrete_ready_mix_3000psi",
             "material_quantity": footprint_sf / 81,  # CY: SF / (27 * thickness_ft)
             "description": "4\" slab on grade with WWM and vapor barrier"},
            {"work_item": "footing_continuous",   "quantity": foundation_lf,
             "material_key": "concrete_ready_mix_3000psi",
             "material_quantity": foundation_lf * 0.37,  # CY per LF for 2'x1' footing
             "description": "Continuous footings 2'-0\" x 1'-0\""},
            {"work_item": "foundation_waterproofing", "quantity": foundation_lf * 1.5,
             "material_key": "foundation_wp_gcp_preprufe160",
             "material_quantity": foundation_lf * 1.5,
             "description": "GCP Preprufe 160 foundation waterproofing"},
        ],
        "Masonry": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "CMU / brick / stone masonry",
             "override_amount": round(total_sf * 25.00),
             "note": (
                 f"MASONRY PLACEHOLDER: $25/SF \u00d7 {total_sf:,.0f} SF = "
                 f"${round(total_sf * 25.00):,.0f} | "
                 f"Includes CMU block, brick veneer, cast stone. "
                 f"Range: $15-45/SF depending on masonry type. "
                 f"GET MASONRY SUB QUOTE."
             )},
        ],
        "Rough Carpentry": [
            {"work_item": "wall_framing_2x6_ext", "quantity": ext_wall_sf,
             "material_key": "lumber_2x6_kd",
             "material_quantity": ext_wall_sf * 1.1,  # LF lumber per SF wall
             "description": "2x6 exterior wall framing at 16\" OC"},
            {"work_item": "wall_framing_2x4_int", "quantity": int_wall_sf,
             "material_key": "lumber_2x4_kd",
             "material_quantity": int_wall_sf * 1.0,
             "description": "2x4 interior wall framing at 16\" OC"},
            {"work_item": "roof_framing_12_12",   "quantity": footprint_sf,
             "material_key": "lumber_2x10_kd",
             "material_quantity": footprint_sf * 0.8,
             "description": "Roof framing 12/12 pitch \u2014 2x10 rafters"},
            {"work_item": "roof_sheathing",       "quantity": roof_sq * 100,
             "material_key": "plywood_5_8_cdx_roof",
             "material_quantity": roof_sq * 100,
             "description": "5/8\" CDX roof sheathing"},
            {"work_item": "wall_sheathing",       "quantity": ext_wall_sf,
             "material_key": "plywood_1_2_cdx_wall",
             "material_quantity": ext_wall_sf,
             "description": "1/2\" CDX wall sheathing"},
        ],
        "Roofing": (
            # EPDM/TPO flat roof for multifamily/commercial
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": (
                     f"EPDM flat roof system: "
                     f"{footprint_sf / 100:.1f} squares"
                 ),
                 "override_amount": round(footprint_sf * 4.50),
                 "note": (
                     f"\u26a0\ufe0f COMMERCIAL ROOFING: EPDM flat roof | "
                     f"${4.50}/SF \u00d7 {footprint_sf:,.0f} SF = "
                     f"${round(footprint_sf * 4.50):,.0f} | "
                     f"Includes membrane, insulation board, flashing, edge metal. "
                     f"Range: $3.50-6.00/SF for EPDM/TPO. "
                     f"GET ROOFING SUB QUOTE."
                 )},
            ] if (is_multifamily or floor_count >= 3) else
            # Residential asphalt shingle (existing logic)
            [
                {"work_item": "shingle_install",      "quantity": roof_sq,
                 "material_key": "shingles_certainteed_xt25",
                 "material_quantity": roof_sq,
                 "description": "CertainTeed XT-25 Nickel Gray shingles"},
                {"work_item": "ice_water_shield_install", "quantity": roof_sq * 0.25,
                 "material_key": "ice_water_shield_grace",
                 "material_quantity": roof_sq * 0.25,
                 "description": "Grace Ice & Water Shield at rakes/ridges/eaves"},
                {"work_item": "felt_install",         "quantity": roof_sq * 0.75,
                 "material_key": "felt_underlayment_30lb",
                 "material_quantity": roof_sq * 0.75,
                 "description": "30 lb felt underlayment"},
                {"work_item": "drip_edge_install",    "quantity": perimeter_lf,
                 "material_key": "drip_edge_aluminum",
                 "material_quantity": perimeter_lf,
                 "description": "Aluminum drip edge"},
                {"work_item": "gutter_install",       "quantity": perimeter_lf,
                 "material_key": "gutter_ogee_5in",
                 "material_quantity": perimeter_lf,
                 "description": "5\" ogee gutters with leaf screens"},
            ]
        ),
        "Siding": [
            {"work_item": "siding_install",       "quantity": ext_wall_sf,
             "material_key": "siding_certainteed_mainstreet_d4",
             "material_quantity": ext_wall_sf * 0.75,
             "description": "CertainTeed Mainstreet Double 4\" Woodgrain"},
            {"work_item": "siding_install",       "quantity": ext_wall_sf * 0.25,
             "material_key": "siding_certainteed_cedar_scallop",
             "material_quantity": ext_wall_sf * 0.25,
             "description": "CertainTeed Cedar Impressions Scallop Colonial White"},
            {"work_item": "housewrap_install",    "quantity": ext_wall_sf,
             "material_key": "housewrap_hydrogap",
             "material_quantity": ext_wall_sf,
             "description": "Benjamin Obdyke HydroGap drainable housewrap"},
            {"work_item": "pvc_trim_install",     "quantity": perimeter_lf * 3,
             "material_key": "pvc_trim_3_4x5_5",
             "material_quantity": perimeter_lf * 3,
             "description": "5-1/2\" PVC trim at windows, doors, corners"},
        ],
        "Insulation": [
            {"work_item": None,                   "quantity": 0,
             "material_key": "spray_foam_closed_cell_r30",
             "material_quantity": ext_wall_sf,
             "description": "R-30 closed cell spray foam exterior walls (sub labor-inclusive)"},
            {"work_item": None,                   "quantity": 0,
             "material_key": "spray_foam_closed_cell_r60",
             "material_quantity": roof_sq * 100,
             "description": "R-60 closed cell spray foam roof assembly (sub labor-inclusive)"},
            {"work_item": "rigid_insulation_foundation", "quantity": foundation_lf * 4,
             "material_key": "rigid_insulation_r10_2in",
             "material_quantity": foundation_lf * 4,
             "description": "R-10 rigid insulation at foundation"},
        ],
        "Drywall": [
            {"work_item": "gwb_hang",             "quantity": gwb_sf,
             "material_key": "gwb_5_8_type_x",
             "material_quantity": gwb_sf,
             "description": "5/8\" Type X GWB hang"},
            {"work_item": "gwb_tape_finish",      "quantity": gwb_sf,
             "material_key": "joint_compound_tape",
             "material_quantity": gwb_sf,
             "description": "Tape, finish, corner bead"},
        ],
        "Flooring": (
            [
                {"work_item": "lvt_install",          "quantity": lvt_sf,
                 "material_key": "lvt_mid_range",
                 "material_quantity": lvt_sf,
                 "description": "LVT-1 Luxury Vinyl Tile all rooms"},
                {"work_item": "tile_install_floor",   "quantity": tile_sf * 0.6,
                 "material_key": "ceramic_tile_floor",
                 "material_quantity": tile_sf * 0.6,
                 "description": "TL-1 Ceramic tile 4x14 bathroom floor"},
                {"work_item": "tile_install_wall_shower", "quantity": tile_sf * 0.4,
                 "material_key": "ceramic_tile_floor",
                 "material_quantity": tile_sf * 0.4,
                 "description": "TL-1 Ceramic tile bathroom wall/shower"},
            ] if is_residential_units else
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": "Flooring allowance",
                 "override_amount": round(total_sf * 12.00),
                 "note": (
                     f"Flooring $12.00/SF \u00d7 {total_sf:,.0f} SF = "
                     f"${round(total_sf * 12.00):,.0f} | "
                     f"Includes LVT, tile, carpet as applicable. "
                     f"GET FLOORING SUB QUOTE."
                 )},
            ]
        ),
        "Painting": [
            {"work_item": "paint_interior_2coat", "quantity": gwb_sf,
             "material_key": "paint_sw_interior_2coat",
             "material_quantity": gwb_sf,
             "description": "SW Alabaster primer + 2 coats interior"},
            {"work_item": "paint_exterior_trim",  "quantity": perimeter_lf * 3,
             "material_key": "paint_exterior_trim",
             "material_quantity": perimeter_lf * 3,
             "description": "Exterior PVC trim paint"},
        ],
        "Plumbing": (
            # Residential multifamily distribution model
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": (
                     f"Commercial plumbing: {max(1, unit_count // 4)} riser stacks + "
                     f"{unit_count} unit rough-ins + {plumbing_fixtures} fixtures"
                 ),
                 "override_amount": round(
                     # Riser stacks
                     max(1, unit_count // 4) * 8500 +
                     # Unit rough-in (supply + drain per unit)
                     unit_count * 4200 +
                     # Fixture allowance
                     plumbing_fixtures * 350 +
                     # Commercial space allowance (if mixed use)
                     (total_sf * 0.20 * 12 if is_commercial else 0)
                 ),
                 "note": (
                     f"\u26a0\ufe0f COMMERCIAL PLUMBING: "
                     f"{max(1, unit_count // 4)} riser stacks + "
                     f"{unit_count} unit rough-ins + "
                     f"{plumbing_fixtures} fixtures. "
                     f"Includes distribution, waste/vent stacks, water heater. "
                     f"GET PLUMBING SUB QUOTE."
                 )},
            ] if (is_residential_units and total_sf > 4000) else
            # Non-residential or small residential — simple $/SF
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": "Plumbing allowance",
                 "override_amount": round(total_sf * 10.00),
                 "note": (
                     f"Plumbing $10.00/SF \u00d7 {total_sf:,.0f} SF = "
                     f"${round(total_sf * 10.00):,.0f} | "
                     f"Includes fixtures, piping, water heater. "
                     f"GET PLUMBING SUB QUOTE."
                 )},
            ] if (is_multifamily and not is_residential_units) else
            # Residential per-fixture model (existing logic)
            [
                {"work_item": "plumbing_rough_per_fixture", "quantity": plumbing_fixtures,
                 "material_key": "plumbing_rough_pipe_per_fix",
                 "material_quantity": plumbing_fixtures,
                 "description": "Supply + drain rough-in per fixture"},
                {"work_item": "plumbing_fixture_set", "quantity": 1,
                 "material_key": "toilet_penguin_254",
                 "material_quantity": 1,
                 "description": "Penguin 254 toilet per spec"},
                {"work_item": "plumbing_fixture_set", "quantity": 1,
                 "material_key": "tub_sterling_ensemble_ada",
                 "material_quantity": 1,
                 "description": "Sterling Ensemble ADA tub per spec"},
                {"work_item": "plumbing_fixture_set", "quantity": 1,
                 "material_key": "shower_set_miseno_mia",
                 "material_quantity": 1,
                 "description": "Miseno Mia shower set per spec"},
            ]
        ),
        "Electrical": [
            {"work_item": "panel_install_200amp", "quantity": 1,
             "material_key": "panel_200amp",
             "material_quantity": 1,
             "description": "200A electrical panel"},
            {"work_item": "electrical_rough_per_circuit", "quantity": electrical_circuits,
             "material_key": "wire_per_circuit_rough",
             "material_quantity": electrical_circuits,
             "description": "Branch circuit rough-in wiring"},
            {"work_item": "device_install",       "quantity": 45,
             "material_key": "outlet_device",
             "material_quantity": 30,
             "description": "Outlets, switches, GFI devices"},
            {"work_item": "light_fixture_recessed", "quantity": 12,
             "material_key": "light_recessed_lithonia_wf6",
             "material_quantity": 12,
             "description": "Lithonia WF6 LED recessed lights"},
        ],
        "HVAC": (
            # Commercial/institutional RTU + ductwork system
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": (
                     f"Commercial HVAC: "
                     f"{max(1, round(total_sf / 4000))} RTU units + "
                     f"supply/return ductwork distribution"
                 ),
                 "override_amount": round(total_sf * (45 if is_residential_units else 32)),
                 "note": (
                     f"\u26a0\ufe0f COMMERCIAL HVAC: "
                     f"${45 if is_residential_units else 32}/SF \u00d7 {total_sf:,.0f} SF = "
                     f"${round(total_sf * (45 if is_residential_units else 32)):,.0f} | "
                     f"Includes RTU units, ductwork, controls, startup. "
                     f"{'Residential multifamily range: $35-55/SF.' if is_residential_units else 'Institutional/commercial range: $25-40/SF.'} "
                     f"RTU vs VRF vs fan coil significantly affects cost. "
                     f"GET HVAC SUB QUOTE."
                 )},
            ] if is_multifamily else
            # Residential mini-split system (existing logic)
            [
                {"work_item": "mini_split_zone_install", "quantity": hvac_zones,
                 "material_key": "mitsubishi_mini_split_zone",
                 "material_quantity": hvac_zones,
                 "description": "Mitsubishi mini-split zones"},
            ]
        ),
        "Cabinets": [
            {"work_item": "gwb_hang",             "quantity": 0,
             "material_key": "cabinets_express_rta_per_ln",
             "material_quantity": cabinet_lf,
             "description": "Express Kitchens Barcelona White RTA cabinets"},
            {"work_item": "gwb_hang",             "quantity": 0,
             "material_key": "vanity_36in_ada",
             "material_quantity": 1,
             "description": "36\" ADA vanity per spec"},
        ],
        "Windows": [
            {"work_item": "window_install",       "quantity": window_count,
             "material_key": "window_marvin_ultimate_dh_med",
             "material_quantity": window_count,
             "description": (
                 f"{'Multifamily' if is_multifamily else 'Residential'} windows "
                 f"— {window_count} EA"
                 f"{' | Note: multifamily windows often bid material-only. Verify install is included.' if is_multifamily else ''}"
             )},
        ],
        "Fire Sprinkler": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "Fire sprinkler system per NFPA 13",
             "override_amount": round(
                 total_sf * (5.50 if is_multifamily else 3.25)
             ),
             "note": (
                 f"\u26a0\ufe0f FIRE SPRINKLER: "
                 f"{'NFPA 13 commercial' if is_multifamily else 'NFPA 13R residential'} "
                 f"wet pipe system | "
                 f"${5.50 if is_multifamily else 3.25}/SF \u00d7 {total_sf:,.0f} SF | "
                 f"Range: $4.50-6.50/SF commercial, $2.50-4.00/SF residential. "
                 f"GET FIRE PROTECTION SUB QUOTE."
             )},
        ],
        "Doors/Hdwr/Finish Carp": (
            [
                {"work_item": "door_install_exterior", "quantity": ext_door_count,
                 "material_key": "door_exterior_insulated_alum",
                 "material_quantity": ext_door_count,
                 "description": "Insulated aluminum exterior doors"},
                {"work_item": "door_install_exterior", "quantity": int_door_count,
                 "material_key": "door_interior_hollow_core",
                 "material_quantity": int_door_count,
                 "description": "Interior hollow core doors"},
            ] if is_residential_units else
            [
                {"work_item": None, "quantity": 0,
                 "material_key": None, "material_quantity": 0,
                 "description": "Door allowance — storefront/interior",
                 "override_amount": round(total_sf * 4.50),
                 "note": (
                     f"Door allowance $4.50/SF \u00d7 {total_sf:,.0f} SF = "
                     f"${round(total_sf * 4.50):,.0f} | "
                     f"Verify storefront scope."
                 )},
            ]
        ),
        "Gutters": [
            {"work_item": "gutter_install", "quantity": 120,
             "material_key": "gutter_ogee_5in", "material_quantity": 120,
             "description": "Ogee gutters 5\" with leaf screens + downspouts"},
        ],
        "SITE WORK / CIVIL": [{
            "work_item": None, "quantity": 0,
            "material_key": None, "material_quantity": 0,
            "description": "PLACEHOLDER \u2014 Civil scope requires sub quotes",
            "override_amount": 160000,
            "note": (
                "\u26a0\ufe0f SITEWORK: $160,000 baseline. "
                "Includes grading, utilities, paving, landscaping. "
                "GET CIVIL SUB QUOTE BEFORE BID."
            ),
        }],
        "Conveying Equipment": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "Ceiling lift system BR5 \u2014 structural + track install",
             "override_amount": 15000,
             "note": "\u26a0\ufe0f Ceiling lift placeholder $15,000. "
                     "Midstate supplies equipment, GC installs. "
                     "Range: $8,000\u2013$25,000. Get specialty sub quote."}
        ],
        "Specialties": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "ADA bathroom accessories \u2014 grab bars, mirrors, hardware",
             "override_amount": 4500,
             "note": "Bobrick grab bars, Harney accessories, ADA mirror. "
                     "Per spec A1.02 equipment plan."}
        ],
        "Countertops & Finishes": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "Kitchen + bath countertops CT-1 Dolomite finish",
             "override_amount": 5500,
             "note": "CT-1 countertop Dolomite 12 UL Dick, kitchen + bath. "
                     "Placeholder \u2014 verify LF with sub."}
        ],
        "Smart Home / Security": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "ARMS system rough-in, smart locks, sensors, doorbell",
             "override_amount": 12000,
             "note": "\u26a0\ufe0f ADA Smart Home scope: ARMS panel, motion/door/flood "
                     "sensors, smart locks, video doorbell, camera. "
                     "Rough-in labor + materials. Range: $8,000\u2013$18,000."}
        ],
        "Structural Steel": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "Simpson hardware, LVL beams, anchor bolts, metal connectors",
             "override_amount": 6500,
             "note": "Simpson H2.5A clips, LSTA24 straps, anchor bolts, "
                     "(2) 1-3/4x11-7/8 LVL ridge beam. Per S1.01."}
        ],
        "Flashing & Waterproofing": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "WR Grace Perm-A-Barrier, aluminum flashing, termite shield",
             "override_amount": 4200,
             "note": "WR Grace Perm-A-Barrier membrane flashing, .060 aluminum "
                     "flashing, termite shield. Per A4.01 details."}
        ],
        "General Requirements": [
            {"work_item": None, "quantity": 0,
             "material_key": None, "material_quantity": 0,
             "description": "General conditions, supervision, temp facilities",
             "override_amount": None,  # calculated dynamically as % of direct cost
             "note": "Calculated as % of direct building cost. "
                     "Includes: superintendent, temp power, dumpsters, "
                     "survey, project signage, insurance, small tools."}
        ],
    }

    return TRADE_MAP.get(trade_name, [])


def _detect_site_scope(consolidated: list[dict], trades: list) -> dict:
    """
    Detect whether site drawings / civil scope are present in the extraction.

    Returns:
    {
        "has_site_scope": bool,
        "site_trade": dict or None,  # the consolidated row if found
        "flag_note": str
    }
    """
    SITE_KEYWORDS = [
        "sitework", "earthwork", "grading", "excavat", "demolit",
        "civil", "d-1", "c-1", "landscape", "dot", "paving",
        "ramp", "curb", "utility", "stormwater", "drainage"
    ]

    # Check if any consolidated trade is already site-related
    site_trade = None
    for row in consolidated:
        lower_trade = row["trade"].lower()
        if any(kw in lower_trade for kw in ["sitework", "paving", "landscaping", "fencing", "site concrete"]):
            site_trade = row
            break

    # Also scan raw trades for civil keywords
    raw_text = " ".join([t.trade.lower() + " " + t.scope_description.lower() for t in trades])
    has_keywords = any(kw in raw_text for kw in SITE_KEYWORDS)

    if site_trade:
        return {
            "has_site_scope": True,
            "site_trade": site_trade,
            "flag_note": ""
        }
    elif has_keywords:
        return {
            "has_site_scope": True,
            "site_trade": None,
            "flag_note": "\u26a0\ufe0f SITE SCOPE DETECTED \u2014 quantities not extracted. "
                         "Estimator must price. Typical range: $25,000\u2013$200,000 depending on scope."
        }
    else:
        return {
            "has_site_scope": False,
            "site_trade": None,
            "flag_note": "\u26a0\ufe0f NO SITE DRAWINGS PROVIDED \u2014 if site work exists, this line "
                         "must be priced separately. Do not bid without site scope."
        }


# ─── Page processing ────────────────────────────────────────────────────────


def _call_gemini_vision(
    png_bytes: bytes,
    prompt_text: str,
    gemini_api_key: str,
    pdf_name: str = "",
    page_num: int = 0,
    emit: Callable = None,
) -> str:
    """
    Send page image to Gemini 3.1 Pro for extraction.
    Returns raw text response matching Claude format.
    Falls back to Claude if Gemini fails.
    """
    import PIL.Image
    import io as _io

    _emit = emit or (lambda x: None)

    try:
        model = _get_gemini_client(gemini_api_key)
        image = PIL.Image.open(_io.BytesIO(png_bytes))

        response = model.generate_content(
            [image, prompt_text],
            request_options={"timeout": 120},
            generation_config=genai.GenerationConfig(
                temperature=0.1,
                max_output_tokens=4096,
            ),
            media_resolution="media_resolution_high",
        )

        raw_text = response.text

        # Strip markdown fences if Gemini wraps JSON
        if raw_text.strip().startswith("```"):
            raw_text = raw_text.strip()
            raw_text = raw_text.split("\n", 1)[1]
            raw_text = raw_text.rsplit("```", 1)[0]

        _emit(f"[GEMINI] {pdf_name} p.{page_num} — "
              f"{len(png_bytes)//1024}KB")

        return raw_text.strip()

    except Exception as e:
        _emit(f"[GEMINI FALLBACK] {pdf_name} p.{page_num}"
              f" — Gemini failed ({e}), using Claude")
        return None  # caller handles fallback


def process_page(
    client: anthropic.Anthropic,
    image_bytes: bytes,
    pdf_name: str,
    page_num: int,
    stats: dict,
    gemini_api_key: str = "",
) -> Optional[PageExtraction]:
    """Send a page image to Claude and parse the extraction result."""

    b64_image = base64.standard_b64encode(image_bytes).decode("utf-8")
    b64_kb = len(b64_image) // 1024
    print(f"[VISION] {pdf_name} p.{page_num} — image size: {len(image_bytes)//1024}KB raw, {b64_kb}KB base64", flush=True)
    prompt_text = EXTRACTION_PROMPT.format(pdf_name=pdf_name, page_num=page_num)

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/png",
                        "data": b64_image,
                    },
                },
                {
                    "type": "text",
                    "text": prompt_text,
                },
            ],
        }
    ]

    SIMPLIFIED_PROMPT = """
This is a construction drawing page.
Extract ANY trades or work items visible.
Return valid JSON only:
{
  "trades_and_scope": [
    {"trade": "trade name", "scope": ["scope item"]}
  ],
  "submission_requirements": [],
  "flags": []
}
If truly nothing is extractable, return:
{"trades_and_scope": [], "submission_requirements": [], "flags": []}
"""

    for attempt in range(4):
        try:
            # On attempts 3 and 4, use simplified prompt
            if attempt >= 2:
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": b64_image,
                                },
                            },
                            {
                                "type": "text",
                                "text": SIMPLIFIED_PROMPT,
                            },
                        ],
                    }
                ]

            response = client.messages.create(
                model=MODEL,
                max_tokens=4096,
                temperature=TEMPERATURE,
                messages=messages,
            )

            usage = response.usage
            stats["input_tokens"] += usage.input_tokens
            stats["output_tokens"] += usage.output_tokens
            stats["api_calls"] += 1

            raw_text = response.content[0].text
            json_text = extract_json_from_text(raw_text)

            data = json.loads(json_text)
            extraction = PageExtraction(**data)
            return extraction

        except (json.JSONDecodeError, ValidationError) as e:
            print(f"[VISION PARSE] {pdf_name} p.{page_num} attempt {attempt+1}/4: "
                  f"{type(e).__name__}: {e}", flush=True)
            if 'raw_text' in dir():
                print(f"[VISION RAW] first 500 chars: {raw_text[:500]}", flush=True)
            if attempt == 0:
                retry_text = RETRY_PROMPT.format(pdf_name=pdf_name, page_num=page_num)
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": b64_image,
                                },
                            },
                            {
                                "type": "text",
                                "text": prompt_text,
                            },
                        ],
                    },
                    {
                        "role": "assistant",
                        "content": raw_text,
                    },
                    {
                        "role": "user",
                        "content": retry_text,
                    },
                ]
            elif attempt == 3:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")

        except anthropic.APIStatusError as e:
            print(f"[VISION API] {pdf_name} p.{page_num} attempt {attempt+1}/4: "
                  f"HTTP {e.status_code}: {e.message}", flush=True)
            if e.status_code == 429:
                wait = 2 ** attempt
                time.sleep(wait)
                continue
            if attempt < 3:
                time.sleep(5)
            else:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")
        except anthropic.APIError as e:
            print(f"[VISION API] {pdf_name} p.{page_num} attempt {attempt+1}/4: "
                  f"APIError: {e}", flush=True)
            if attempt < 3:
                time.sleep(5)
            else:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")

    # Return empty but valid result — never lose a page silently
    return PageExtraction(
        trades_and_scope=[],
        submission_requirements=[],
        flags=[f"WARNING: Page {page_num} could not be parsed after 4 attempts — manual review required"]
    )


# ─── Orchestrator ────────────────────────────────────────────────────────────


_CAD_KEYWORD_PAT = re.compile(
    r'(?:FLOOR|ROOF|SITE|FOUNDATION)\s*PLAN'
    r'|ELEVATION|SECTION\s+[A-Z]|DETAIL\s+\d'
    r'|SCALE\s*[:=]|NOT\s+TO\s+SCALE'
    r'|[A-Z]\d{1,3}\.\d'           # sheet refs: A1.1, S2.0
    r'|\bTYP\.?\b|\bSIM\.?\b'       # CAD abbreviations
    r'|\bEQ\.?\b|\bN\.T\.S',
    re.IGNORECASE,
)


def _classify_page(doc: fitz.Document, page_idx: int) -> str:
    """Classify a PDF page to choose extraction method.

    Returns:
        'text'           — mostly text, use native text extraction (cheapest)
        'mixed'          — text + images/tables, use Vision API
        'image_drawing'  — drawings/floor plans with embedded images, use Vision API
        'image_scanned'  — scanned page, no extractable text, use Vision API
                           (future: route to OCR tier like Document AI instead)
    """
    page = doc.load_page(page_idx)
    text = page.get_text("text").strip()
    char_count = len(text)

    # Count images on the page
    image_list = page.get_images(full=True)
    image_count = len(image_list)

    # Calculate image area as fraction of page
    page_rect = page.rect
    page_area = page_rect.width * page_rect.height
    image_area = 0
    for img in image_list:
        try:
            bbox = page.get_image_bbox(img)
            if bbox:
                image_area += abs(bbox.width * bbox.height)
        except Exception:
            pass
    image_fraction = image_area / page_area if page_area > 0 else 0

    # Count vector drawing objects (lines, curves, paths from CAD)
    drawings = page.get_drawings()
    vector_count = len(drawings)

    # ── Low/no text pages ──
    if char_count < 50 and image_count > 0:
        return "image_drawing"
    elif char_count < 50 and image_count == 0:
        if vector_count > 50:
            return "image_drawing"   # Vector-only CAD drawing
        return "image_scanned"

    # ── CAD drawing detection (5 heuristic signals) ──
    # CAD-generated PDFs embed text layers (title blocks, dimensions, room
    # labels) that push char_count well over 50.  We score 5 signals and
    # classify as drawing when 3+ fire (or 2+ with images present).
    lines = [ln for ln in text.split('\n') if ln.strip()]
    num_lines = max(len(lines), 1)
    avg_line_len = char_count / num_lines

    signals = 0
    signal_details = []

    # Signal 1: Short average line length (< 22 chars — CAD labels vs paragraphs)
    if avg_line_len < 22:
        signals += 1
        signal_details.append("short_lines")

    # Signal 2: High digit ratio (> 15% — dimensions like 12'-6", 3/4", coords)
    digit_count = sum(1 for c in text if c.isdigit())
    digit_ratio = digit_count / max(char_count, 1)
    if digit_ratio > 0.15:
        signals += 1
        signal_details.append("high_digits")

    # Signal 3: Drawing keywords density (SCALE, DETAIL, SECTION, sheet refs)
    kw_hits = len(_CAD_KEYWORD_PAT.findall(text))
    kw_density = kw_hits / num_lines
    if kw_density > 0.3:
        signals += 1
        signal_details.append("drawing_keywords")

    # Signal 4: Many very short lines (> 50% under 10 chars — room labels, callouts)
    short_lines = sum(1 for ln in lines if len(ln.strip()) < 10)
    short_ratio = short_lines / num_lines
    if short_ratio > 0.5:
        signals += 1
        signal_details.append("short_ratio")

    # Signal 5: High vector drawing object count (> 50 — lines/curves from CAD)
    if vector_count > 50:
        signals += 1
        signal_details.append("vectors")

    # Decision: 3+ signals = CAD drawing.  2+ signals with images also qualifies.
    is_cad = signals >= 3 or (signals >= 2 and image_count > 0)

    if is_cad:
        print(f"[CLASSIFY] p.{page_idx+1} detected as CAD drawing "
              f"(chars={char_count}, avg_line={avg_line_len:.0f}, "
              f"digit_ratio={digit_ratio:.2f}, short_ratio={short_ratio:.2f}, "
              f"vectors={vector_count}, signals={signals}/5 [{', '.join(signal_details)}])",
              flush=True)
        return "image_drawing"

    if image_fraction > 0.4:
        return "mixed"
    else:
        return "text"


def _extract_text_page(
    client: anthropic.Anthropic,
    text: str,
    pdf_name: str,
    page_num: int,
    stats: dict,
) -> Optional[PageExtraction]:
    """Process a text-heavy page using native text extraction (no Vision).
    Sends extracted text to Claude as a text prompt — much cheaper than Vision."""

    prompt_text = EXTRACTION_PROMPT.format(pdf_name=pdf_name, page_num=page_num)

    full_prompt = (
        f"--- DOCUMENT CONTENT ({pdf_name}, Page {page_num}) ---\n\n"
        f"{text}\n\n"
        f"--- END DOCUMENT CONTENT ---\n\n"
        f"{prompt_text}"
    )

    for attempt in range(2):
        try:
            response = client.messages.create(
                model=MODEL,
                max_tokens=4096,
                temperature=TEMPERATURE,
                messages=[{"role": "user", "content": full_prompt}],
            )

            usage = response.usage
            stats["input_tokens"] += usage.input_tokens
            stats["output_tokens"] += usage.output_tokens
            stats["api_calls"] += 1

            raw_text = response.content[0].text
            json_text = extract_json_from_text(raw_text)
            data = json.loads(json_text)
            extraction = PageExtraction(**data)
            return extraction

        except (json.JSONDecodeError, ValidationError) as e:
            print(f"[TEXT PARSE] {pdf_name} p.{page_num} attempt {attempt+1}/2: "
                  f"{type(e).__name__}: {e}", flush=True)
            if 'raw_text' in dir():
                print(f"[TEXT RAW] first 500 chars: {raw_text[:500]}", flush=True)
            if attempt == 1:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")
        except anthropic.APIStatusError as e:
            if e.status_code == 429:
                wait = 2 ** attempt
                time.sleep(wait)
                continue
            if attempt == 0:
                time.sleep(3)
            else:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")
        except anthropic.APIError:
            if attempt == 0:
                time.sleep(3)
            else:
                stats["failed_pages"].append(f"{pdf_name} p.{page_num}")

    return PageExtraction(
        trades_and_scope=[],
        submission_requirements=[],
        flags=[f"WARNING: Page {page_num} text extraction failed — manual review required"]
    )


def _process_pdf(client, file_path, emit, global_stats, file_stats, gemini_api_key=""):
    """Process a PDF file page by page using parallel extraction (8 workers)."""
    file_name = file_path.name
    file_path_str = str(file_path)
    doc = fitz.open(file_path_str)
    num_pages = len(doc)
    emit(f"{file_name} has {num_pages} pages")

    per_file = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0,
                "pages": num_pages, "failed_pages": []}
    reqs, trades = [], []

    # Classify all pages up front (read-only, safe on main thread)
    page_types = {}
    for page_idx in range(num_pages):
        page_types[page_idx] = _classify_page(doc, page_idx)
        per_file[f"{page_types[page_idx]}_pages"] = \
            per_file.get(f"{page_types[page_idx]}_pages", 0) + 1

    # Classification summary
    n_text = sum(1 for t in page_types.values() if t == "text")
    n_draw = sum(1 for t in page_types.values() if t in ("image_drawing", "mixed"))
    n_scan = sum(1 for t in page_types.values() if t == "image_scanned")
    emit(f"  Classified: {n_text} text, {n_draw} drawing/mixed, {n_scan} scanned")

    # Pre-extract text for text pages (read-only, safe on main thread)
    page_texts = {}
    for page_idx, ptype in page_types.items():
        if ptype == "text":
            page_texts[page_idx] = doc.load_page(page_idx).get_text("text").strip()

    doc.close()  # Close main doc — workers open their own for rendering

    MAX_WORKERS = 8
    results = {}  # page_idx → extraction
    lock = threading.Lock()

    def process_single_page(page_idx):
        page_num = page_idx + 1
        page_type = page_types[page_idx]
        page_stats = {
            "input_tokens": 0, "output_tokens": 0,
            "api_calls": 0, "failed_pages": []
        }

        if page_type == "text":
            # Tier 1: Native text extraction — cheapest
            emit(f"Processing {file_name} p.{page_num}/{num_pages} [TEXT]...")
            extraction = _extract_text_page(
                client, page_texts[page_idx], file_name, page_num, page_stats
            )
            # Fallback: if text extraction failed, escalate to Vision
            text_failed = (
                (extraction and extraction.flags and not extraction.trades_and_scope)
                or any(f"p.{page_num}" in fp for fp in page_stats["failed_pages"])
            )
            if text_failed:
                emit(f"  \u21b3 Text extraction failed for p.{page_num} \u2014 retrying with Vision...")
                try:
                    thread_doc = fitz.open(file_path_str)
                    png_bytes = render_page_to_png(thread_doc, page_idx)
                    thread_doc.close()
                    vision_stats = {"input_tokens": 0, "output_tokens": 0,
                                    "api_calls": 0, "failed_pages": []}
                    vision_extraction = process_page(
                        client, png_bytes, file_name, page_num, vision_stats,
                        gemini_api_key=gemini_api_key,
                    )
                    # Accumulate Vision stats regardless
                    page_stats["input_tokens"] += vision_stats["input_tokens"]
                    page_stats["output_tokens"] += vision_stats["output_tokens"]
                    page_stats["api_calls"] += vision_stats["api_calls"]

                    if vision_extraction and not vision_extraction.flags:
                        extraction = vision_extraction  # Vision succeeded
                        # Clear text-attempt failure entries
                        page_stats["failed_pages"] = [
                            fp for fp in page_stats["failed_pages"]
                            if f"p.{page_num}" not in fp
                        ]
                    else:
                        # Vision also failed — keep original failure
                        page_stats["failed_pages"].extend(vision_stats["failed_pages"])
                except Exception:
                    pass  # Keep the original text extraction result
        else:
            # Tier 2/3: Vision API — each worker opens its own fitz doc for rendering
            label = "SCAN" if page_type == "image_scanned" else page_type.upper().replace("IMAGE_", "")
            emit(f"Processing {file_name} p.{page_num}/{num_pages} [{label}]...")
            try:
                thread_doc = fitz.open(file_path_str)
                png_bytes = render_page_to_png(thread_doc, page_idx)
                thread_doc.close()
            except Exception as e:
                emit(f"[RENDER ERROR] {file_name} p.{page_num}: {e}")
                with lock:
                    per_file["failed_pages"].append(f"{file_name} p.{page_num}")
                return page_idx, None

            extraction = process_page(
                client, png_bytes, file_name, page_num, page_stats,
                gemini_api_key=gemini_api_key,
            )

        # Thread-safe stats accumulation
        with lock:
            per_file["input_tokens"] += page_stats["input_tokens"]
            per_file["output_tokens"] += page_stats["output_tokens"]
            per_file["api_calls"] += page_stats["api_calls"]
            per_file["failed_pages"].extend(page_stats["failed_pages"])

        return page_idx, extraction

    # Run pages in parallel
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {
            executor.submit(process_single_page, i): i
            for i in range(num_pages)
        }
        completed = 0
        for future in as_completed(futures):
            completed += 1
            try:
                page_idx, extraction = future.result()
                if extraction is not None:
                    results[page_idx] = extraction
                if completed % 10 == 0 or completed == num_pages:
                    emit(f"  \u23f3 {file_name}: {completed}/{num_pages} pages done...")
            except Exception as e:
                emit(f"[PAGE ERROR] {e}")

    # Process results IN ORDER (page 1 before page 2)
    for page_idx in sorted(results.keys()):
        extraction = results[page_idx]
        page_num = page_idx + 1

        if extraction.flags:
            emit(f"\u26a0\ufe0f {file_name} p.{page_num} PARSE FAILED \u2014 flagged for manual review")
        else:
            emit(f"{file_name} p.{page_num} OK (reqs={len(extraction.submission_requirements)}, trades={len(extraction.trades_and_scope)})")
        reqs.extend(extraction.submission_requirements)
        trades.extend(extraction.trades_and_scope)

    file_stats[file_name] = per_file
    for k in ["input_tokens", "output_tokens", "api_calls"]:
        global_stats[k] += per_file[k]
    global_stats["total_pages"] += per_file["pages"]
    global_stats["failed_pages"].extend(per_file["failed_pages"])

    _tp = per_file.get('text_pages', 0)
    _mp = per_file.get('mixed_pages', 0)
    _id = per_file.get('image_drawing_pages', 0)
    _is = per_file.get('image_scanned_pages', 0)
    emit(f"  \U0001f4ca {file_name}: {_tp} text / {_mp} mixed / {_id} drawing / {_is} scanned pages "
         f"(Vision calls saved: {_tp})")
    # Track scanned pages — when this number is consistently high across projects,
    # adding a dedicated OCR tier (Google Document AI @ $1.50/1000 pages) will
    # replace Vision for these pages and cut cost further.
    if _is > 5:
        emit(f"  \U0001f4a1 {file_name}: {_is} scanned pages hit Vision API — "
             f"OCR tier (Document AI) would save ~${_is * 0.05:.2f} on this file")

    return reqs, trades


def _process_image(client, file_path, emit, global_stats, file_stats, gemini_api_key=""):
    """Process a standalone image file via Vision API."""
    file_name = file_path.name
    emit(f"Processing image {file_name}...")

    image_bytes = file_path.read_bytes()
    if len(image_bytes) > MAX_IMAGE_BYTES:
        emit(f"[WARNING] {file_name} is large ({len(image_bytes)} bytes), may be truncated")

    media_type = _get_image_media_type(file_path)
    per_file = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0,
                "pages": 1, "failed_pages": []}

    # Use the same process_page but with the raw image bytes
    page_stats = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0, "failed_pages": []}

    b64_image = base64.standard_b64encode(image_bytes).decode("utf-8")
    prompt_text = EXTRACTION_PROMPT.format(pdf_name=file_name, page_num=1)

    messages = [{"role": "user", "content": [
        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_image}},
        {"type": "text", "text": prompt_text},
    ]}]

    SIMPLIFIED_PROMPT = """
This is a construction drawing image.
Extract ANY trades or work items visible.
Return valid JSON only:
{
  "trades_and_scope": [
    {"trade": "trade name", "scope": ["scope item"]}
  ],
  "submission_requirements": [],
  "flags": []
}
If truly nothing is extractable, return:
{"trades_and_scope": [], "submission_requirements": [], "flags": []}
"""

    extraction = None
    for attempt in range(4):
        try:
            # On attempts 3 and 4, use simplified prompt
            if attempt >= 2:
                messages = [{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_image}},
                    {"type": "text", "text": SIMPLIFIED_PROMPT},
                ]}]

            response = client.messages.create(model=MODEL, max_tokens=4096,
                                              temperature=TEMPERATURE, messages=messages)
            page_stats["input_tokens"] += response.usage.input_tokens
            page_stats["output_tokens"] += response.usage.output_tokens
            page_stats["api_calls"] += 1
            raw_text = response.content[0].text
            data = json.loads(extract_json_from_text(raw_text))
            extraction = PageExtraction(**data)
            break
        except (json.JSONDecodeError, ValidationError):
            if attempt == 0:
                messages.append({"role": "assistant", "content": raw_text})
                messages.append({"role": "user", "content": RETRY_PROMPT.format(pdf_name=file_name, page_num=1)})
            elif attempt == 3:
                page_stats["failed_pages"].append(file_name)
        except anthropic.APIError:
            if attempt < 3:
                time.sleep(5)
            else:
                page_stats["failed_pages"].append(file_name)

    if extraction is None:
        extraction = PageExtraction(
            trades_and_scope=[],
            submission_requirements=[],
            flags=[f"WARNING: {file_name} could not be parsed after 4 attempts — manual review required"]
        )

    per_file["input_tokens"] = page_stats["input_tokens"]
    per_file["output_tokens"] = page_stats["output_tokens"]
    per_file["api_calls"] = page_stats["api_calls"]
    per_file["failed_pages"] = page_stats["failed_pages"]
    file_stats[file_name] = per_file
    for k in ["input_tokens", "output_tokens", "api_calls"]:
        global_stats[k] += per_file[k]
    global_stats["total_pages"] += 1
    global_stats["failed_pages"].extend(per_file["failed_pages"])

    reqs, trades = [], []
    reqs = list(extraction.submission_requirements)
    trades = list(extraction.trades_and_scope)
    if extraction.flags:
        emit(f"\u26a0\ufe0f {file_name} PARSE FAILED \u2014 flagged for manual review")
    else:
        emit(f"{file_name} OK (reqs={len(reqs)}, trades={len(trades)})")
    return reqs, trades


def _process_text_file(client, file_path, emit, global_stats, file_stats):
    """Process a text-based file (DOCX, EML, XLSX, CSV, TXT) by extracting text and sending to Claude."""
    file_name = file_path.name
    ext = file_path.suffix.lower()

    emit(f"Extracting text from {file_name}...")

    try:
        if ext == ".docx":
            text = _extract_text_from_docx(file_path)
        elif ext == ".eml":
            text = _extract_text_from_eml(file_path)
        elif ext in (".xlsx", ".xls", ".csv"):
            text = _extract_text_from_spreadsheet(file_path)
        elif ext in (".txt", ".rtf", ".doc"):
            text = _extract_text_from_txt(file_path)
        else:
            text = _extract_text_from_txt(file_path)
    except Exception as e:
        emit(f"[ERROR] Could not read {file_name}: {e}")
        return [], []

    if not text.strip():
        emit(f"[WARNING] {file_name} has no extractable text, skipping")
        return [], []

    # Chunk the text and process each chunk
    chunks = _chunk_text(text)
    num_chunks = len(chunks)
    emit(f"{file_name}: {len(text)} chars, {num_chunks} section(s)")

    per_file = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0,
                "pages": num_chunks, "failed_pages": []}
    reqs, trades = [], []

    for chunk_idx, chunk in enumerate(chunks):
        section_num = chunk_idx + 1
        emit(f"Processing {file_name} — Section {section_num}/{num_chunks}...")

        page_stats = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0, "failed_pages": []}
        extraction = process_text_page(client, chunk, file_name, section_num, page_stats)

        per_file["input_tokens"] += page_stats["input_tokens"]
        per_file["output_tokens"] += page_stats["output_tokens"]
        per_file["api_calls"] += page_stats["api_calls"]
        per_file["failed_pages"].extend(page_stats["failed_pages"])

        if extraction.flags:
            emit(f"\u26a0\ufe0f {file_name} s.{section_num} PARSE FAILED \u2014 flagged for manual review")
        else:
            emit(f"{file_name} s.{section_num} OK (reqs={len(extraction.submission_requirements)}, trades={len(extraction.trades_and_scope)})")
        reqs.extend(extraction.submission_requirements)
        trades.extend(extraction.trades_and_scope)

    file_stats[file_name] = per_file
    for k in ["input_tokens", "output_tokens", "api_calls"]:
        global_stats[k] += per_file[k]
    global_stats["total_pages"] += per_file["pages"]
    global_stats["failed_pages"].extend(per_file["failed_pages"])
    return reqs, trades


def classify_project_complexity(project_info: dict, page_texts: list) -> dict:
    """
    Detects construction type from document keywords.
    Adds construction_type, complexity_multiplier, and
    complexity_warning to the project_info dict.
    """
    GUT_REHAB_KEYWORDS = [
        "gut rehabilitation", "gut rehab", "full gut",
        "strip to studs", "abatement", "hazmat",
        "lead paint", "asbestos", "complete gut"
    ]
    RENOVATION_KEYWORDS = [
        "renovation", "rehab", "rehabilitation",
        "existing building", "existing structure",
        "selective demo", "selective demolition",
        "infill", "adaptive reuse", "historic",
        "occupied", "phased construction",
        "alteration", "alterations", "retrofit",
        "upgrade", "modernization"
    ]
    ADDITION_KEYWORDS = [
        "addition", "expansion", "new wing",
        "building addition", "annex"
    ]

    all_text = " ".join(page_texts).lower()
    triggered_keywords = []
    detected_type = project_info.get("construction_type") or "unknown"

    if detected_type in ("unknown", None):
        for kw in GUT_REHAB_KEYWORDS:
            if kw in all_text:
                triggered_keywords.append(kw)
                detected_type = "gut_rehabilitation"
                break

        if detected_type in ("unknown", None):
            for kw in RENOVATION_KEYWORDS:
                if kw in all_text:
                    triggered_keywords.append(kw)
                    detected_type = "renovation"
                    break

        if detected_type in ("unknown", None):
            for kw in ADDITION_KEYWORDS:
                if kw in all_text:
                    triggered_keywords.append(kw)
                    detected_type = "addition"
                    break

    COMPLEXITY_CONFIG = {
        "new_construction": {
            "multiplier": 1.0,
            "warning": None,
        },
        "renovation": {
            "multiplier": 1.35,
            "warning": (
                "\u26a0\ufe0f RENOVATION PROJECT: All costs include 1.35x "
                "renovation premium. Actual costs depend heavily "
                "on existing conditions. Sub quotes required "
                "before bid. Walk the site first."
            ),
        },
        "gut_rehabilitation": {
            "multiplier": 1.55,
            "warning": (
                "\u26a0\ufe0f GUT REHABILITATION: All costs include 1.55x "
                "gut rehab premium. Significant unknowns exist. "
                "Do NOT bid without site walk, sub quotes, and "
                "existing conditions survey."
            ),
        },
        "addition": {
            "multiplier": 1.20,
            "warning": (
                "\u26a0\ufe0f BUILDING ADDITION: All costs include 1.20x "
                "addition premium for tie-in complexity, "
                "phasing, and existing structure interface."
            ),
        },
        "unknown": {
            "multiplier": 1.0,
            "warning": (
                "\u26a0\ufe0f CONSTRUCTION TYPE UNKNOWN: Could not determine "
                "if this is new construction or renovation. "
                "Verify before using this estimate."
            ),
        },
    }

    config = COMPLEXITY_CONFIG.get(detected_type, COMPLEXITY_CONFIG["unknown"])

    project_info["construction_type"] = detected_type
    project_info["complexity_multiplier"] = config["multiplier"]
    project_info["complexity_warning"] = config["warning"]
    project_info["construction_type_keywords"] = triggered_keywords

    return project_info


def extract_project_quantities(trades: list, emit: Callable[[str], None], api_key: str,
                               timeout_seconds: int = 60) -> dict:
    """
    Asks Claude to extract key project dimensions from all extracted trade data.
    Returns a quantities dict that gets passed into get_baseline_quantities().
    Times out after timeout_seconds and returns {} with a warning.
    """
    import concurrent.futures

    def _do_extraction() -> dict:
        return _extract_project_quantities_inner(trades, emit, api_key)

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_do_extraction)
            return future.result(timeout=timeout_seconds)
    except concurrent.futures.TimeoutError:
        emit(
            f"[WARNING] Project quantity extraction timed out after {timeout_seconds}s "
            f"({len(trades)} trades). Using conservative defaults. "
            f"Results may be less accurate for this project."
        )
        return {}
    except Exception as e:
        emit(f"[WARNING] Project quantity extraction failed: {e}. Using conservative defaults.")
        return {}


def _extract_project_quantities_inner(trades: list, emit: Callable[[str], None], api_key: str) -> dict:
    """Inner implementation of extract_project_quantities (no timeout wrapper).

    Uses two sequential API calls:
      Call 1 — Project identity: project_type, construction_type, unit_count, floor_count
      Call 2 — Physical dimensions: uses Call 1 results as context for remaining fields
    This prevents fixture/room counts from being confused with dwelling units.
    """
    # Compile all scope text from extractions
    scope_text = []
    for t in trades[:200]:
        scope_text.append(f"{t.trade}: {t.scope_description} [{t.evidence}]")

    combined = "\n".join(scope_text)
    emit(f"[QUANTITIES] combined text length: {len(combined)} chars, {len(scope_text)} trade entries")
    emit(f"[QUANTITIES SAMPLE] first 300 chars: {combined[:300]}")

    # ── CALL 1: Project identity (focused, unambiguous) ──────────────────────
    prompt_identity = f"""You are analyzing construction documents.
Extract ONLY the 4 fields below. Return ONLY valid JSON, no other text.

{{
  "project_type": "single_family" | "multi_family" | "commercial" | "mixed_use" | "warehouse" | "institutional" | "unknown",
  "construction_type": "new_construction" | "renovation" | "gut_rehabilitation" | "addition" | "unknown",
  "unit_count": <integer — see definition below>,
  "floor_count": <number of stories above grade>
}}

CRITICAL DEFINITION — unit_count:
Count ONLY self-contained residential dwelling units — apartments, condos,
townhomes, or SRO sleeping rooms that are someone's primary residence.
A dwelling unit has its own entrance, bathroom, and sleeping area.

DO NOT COUNT:
- Rooms in a community center, library, office, school, or church
- Plumbing fixtures (toilets, sinks, showers)
- Parking spaces, storage units, or amenity rooms
- Meeting rooms, classrooms, gym spaces, or offices
- Any space that is not a residential dwelling unit

If the building is a community center, library, school, office building,
church, recreation center, or any non-residential use: unit_count = 0.
If single family home: unit_count = 1.
If genuinely uncertain: return 0, not a guess.

construction_type: look for keywords like 'new construction', 'renovation',
'rehabilitation', 'existing building', 'alteration', 'addition'.

If a value cannot be determined, use null — never guess.

CONSTRUCTION DOCUMENTS TEXT:
{combined}
"""

    try:
        client = anthropic.Anthropic(api_key=api_key)

        # Call 1: Project identity
        emit("Extracting project identity (type, units, floors)...")
        resp1 = client.messages.create(
            model=MODEL,
            max_tokens=300,
            messages=[{"role": "user", "content": prompt_identity}]
        )
        text1 = resp1.content[0].text.strip().replace('```json', '').replace('```', '').strip()
        identity = json.loads(text1)

        _pt = identity.get('project_type', 'unknown')
        _ct = identity.get('construction_type', 'unknown')
        _uc = identity.get('unit_count') or 0
        _fc = identity.get('floor_count') or 1
        emit(f"Identity: {_pt}, {_ct}, {_uc} units, {_fc} floors")

        # ── CALL 2: Physical dimensions (uses Call 1 as context) ─────────────
        prompt_dimensions = f"""You are analyzing construction documents.
This is a {_pt} building with {_uc} residential dwelling units and {_fc} floors.
Construction type: {_ct}.

Extract physical dimensions. Return ONLY valid JSON, no other text.

CRITICAL: total_building_sf is the MOST IMPORTANT scalar — every trade estimate
depends on it. Search carefully for Gross Floor Area, Building Area, energy code
tables, zoning tables, or any reference to total square footage.
If footprint_sf and floor_count are both known, COMPUTE total_building_sf =
footprint_sf x floor_count rather than returning null.

{{
  "floor_to_floor_height_ft": <typical floor-to-floor height in feet>,
  "footprint_sf": <building footprint in SF, ground floor only>,
  "total_building_sf": <total gross SF all floors>,
  "perimeter_lf": <building perimeter in LF>,
  "unit_avg_sf": <average dwelling unit size in SF, or null if no units>,
  "roof_type": "epdm_flat" | "tpo_flat" | "asphalt_shingle" | "built_up" | "metal" | "unknown",
  "hvac_system_type": "commercial_rtu" | "residential_minisplit" | "residential_ducted" | "commercial_vrf" | "unknown",
  "foundation_type": "slab_on_grade" | "strip_footing" | "mat_foundation" | "deep_foundation" | "unknown",
  "window_count": <total windows in building>,
  "ext_door_count": <total exterior doors>,
  "int_door_count": <total interior doors in building, count from door schedule if available, else null>,
  "plumbing_fixtures": <total plumbing fixtures in building>,
  "hvac_zones": <number of HVAC zones>,
  "electrical_panels": <number of electrical panels>,
  "circuits_per_unit": <estimated circuits per dwelling unit, or null if no units>,
  "parking_spaces": <if applicable>,
  "elevator_count": <number of elevators>,
  "ada_units": <number of ADA accessible dwelling units, or 0 if no units>,
  "confidence": "high" | "medium" | "low",
  "notes": "brief explanation of how quantities were derived"
}}

Look for system types in scope descriptions, not just cover sheet data.
Examples of what to look for:

roof_type: 'EPDM', 'TPO', 'asphalt shingle', 'CertainTeed', 'built-up',
'flat roof', 'shingle', 'membrane'

hvac_system_type: 'mini-split', 'Mitsubishi', 'Daikin', 'ductless', 'RTU',
'rooftop unit', 'fan coil', 'VRF', 'forced air', 'heat pump', 'PTAC'

foundation_type: 'slab on grade', 'strip footing', 'mat foundation',
'spread footing', 'grade beam', 'pile', 'caisson', 'crawl space'

For footprint_sf: if you find a slab or concrete calculation
in scope descriptions (e.g. 'slab on grade 699 SF' or
'concrete slab 3,115 SF'), use that number as footprint_sf.
If total_building_sf and floor_count are both known,
estimate footprint_sf = total_building_sf / floor_count.

For perimeter_lf: if footprint_sf is known and perimeter_lf
is not explicitly stated, estimate it as:
perimeter_lf = round(sqrt(footprint_sf) * 4)

If a value cannot be determined from the documents, use null — never guess.

CONSTRUCTION DOCUMENTS TEXT:
{combined}
"""

        emit("Extracting physical dimensions...")
        resp2 = client.messages.create(
            model=MODEL,
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt_dimensions}]
        )
        text2 = resp2.content[0].text.strip().replace('```json', '').replace('```', '').strip()
        dimensions = json.loads(text2)

        # Merge: identity fields take priority
        quantities = {**dimensions}
        quantities['project_type'] = _pt
        quantities['construction_type'] = _ct
        quantities['unit_count'] = _uc
        quantities['floor_count'] = _fc

        # Default floor height by project type if not extracted
        if not quantities.get('floor_to_floor_height_ft'):
            _ft_defaults = {
                'single_family': 9, 'multi_family': 9,
                'commercial': 12, 'mixed_use': 12,
                'warehouse': 16, 'institutional': 14,
                'unknown': 9,
            }
            pt = quantities.get('project_type', 'unknown')
            quantities['floor_to_floor_height_ft'] = _ft_defaults.get(pt, 9)
            emit(f"Floor height: {quantities['floor_to_floor_height_ft']}ft (default for {pt})")
        else:
            emit(f"Floor height: {quantities['floor_to_floor_height_ft']}ft (extracted from documents)")

        emit(
            f"Project quantities extracted: "
            f"{quantities.get('project_type')} | "
            f"{quantities.get('unit_count')} units | "
            f"{quantities.get('total_building_sf')} SF | "
            f"confidence: {quantities.get('confidence')}"
        )

        # ── Confidence scoring for scalar gate ────────────────────────────────
        confidence = {}

        key_scalars = [
            'unit_count', 'floor_count', 'total_building_sf', 'footprint_sf',
            'perimeter_lf', 'project_type', 'construction_type',
            'roof_type', 'hvac_system_type', 'foundation_type',
        ]

        for key in key_scalars:
            val = quantities.get(key)
            if val is None or val == 0 or val == 'unknown':
                confidence[key] = 'low'      # Missing — estimator MUST provide
            elif isinstance(val, (int, float)) and val > 0:
                confidence[key] = 'high'     # Extracted, looks reasonable
            else:
                confidence[key] = 'medium'   # Extracted but may need verification

        # Plausibility checks that lower confidence
        _tsf = quantities.get('total_building_sf', 0) or 0
        _uc = quantities.get('unit_count', 0) or 0
        _fc = quantities.get('floor_count', 0) or 0

        if _tsf > 0 and _uc > 0:
            sf_per_unit = _tsf / _uc
            if sf_per_unit < 150 or sf_per_unit > 3000:
                confidence['unit_count'] = 'low'
                confidence['total_building_sf'] = 'low'

        if _tsf > 0 and _fc > 0:
            sf_per_floor = _tsf / _fc
            if sf_per_floor < 500 or sf_per_floor > 50000:
                confidence['floor_count'] = 'low'

        quantities['_confidence'] = confidence
        quantities['_scalar_gate_required'] = any(
            v == 'low' for v in confidence.values()
        )

        return quantities
    except Exception as e:
        emit(f"[WARNING] Could not extract quantities: {e}. Using conservative defaults.")
        return {}


def validate_project_quantities(pq: dict, emit: Callable[[str], None]) -> dict:
    """
    Validate extracted project quantities. Pop any field that fails
    plausibility checks and log a warning. Returns the cleaned dict.
    """
    import math
    findings = []

    def _pop_bad(field, value, reason):
        pq.pop(field, None)
        msg = (f"Extracted {field} = {value} appeared unreasonable ({reason}) "
               f"and was replaced with a conservative default. Please verify.")
        findings.append(msg)
        emit(f"[VALIDATION] {msg}")

    # Auto-estimate total_building_sf from footprint x floors if missing
    if not pq.get('total_building_sf') and pq.get('footprint_sf') and pq.get('floor_count'):
        _est = round(pq['footprint_sf'] * pq['floor_count'])
        pq['total_building_sf'] = _est
        msg = f"Auto-estimated total_building_sf = {pq['footprint_sf']} x {pq['floor_count']} = {_est}"
        findings.append(msg)
        emit(f"[VALIDATION] {msg}")

    # unit_count: 1–999
    uc = pq.get('unit_count')
    if uc is not None and (not isinstance(uc, (int, float)) or uc < 1 or uc > 999):
        _pop_bad('unit_count', uc, "must be 1-999")

    # floor_count: 1–80
    fc = pq.get('floor_count')
    if fc is not None and (not isinstance(fc, (int, float)) or fc < 1 or fc > 80):
        _pop_bad('floor_count', fc, "must be 1-80")

    # footprint_sf: cannot exceed total_building_sf
    fp = pq.get('footprint_sf')
    tsf = pq.get('total_building_sf')
    if fp is not None and tsf is not None:
        if isinstance(fp, (int, float)) and isinstance(tsf, (int, float)) and fp > tsf:
            _pop_bad('footprint_sf', fp, f"exceeds total_building_sf={tsf}")

    # total_building_sf: cross-check with footprint * floors
    if tsf is not None and isinstance(tsf, (int, float)):
        fp_check = pq.get('footprint_sf')
        fc_check = pq.get('floor_count') or 1
        if fp_check and isinstance(fp_check, (int, float)) and fc_check > 1:
            expected = fp_check * fc_check
            if tsf < expected * 0.7 or tsf > expected * 1.3:
                # Replace with corrected value instead of deleting
                corrected = round(expected)
                msg = (f"Extracted total_building_sf = {tsf} inconsistent with "
                       f"footprint={fp_check} x floors={fc_check} = {expected}. "
                       f"Auto-corrected to {corrected}. Please verify.")
                findings.append(msg)
                emit(f"[VALIDATION] {msg}")
                pq['total_building_sf'] = corrected
        else:
            # Standalone check by project type
            pt = pq.get('project_type', 'unknown')
            max_sf = 2_000_000 if pt in ('commercial', 'warehouse', 'institutional', 'mixed_use') else 500_000
            if tsf > max_sf:
                _pop_bad('total_building_sf', tsf, f"exceeds {max_sf:,} for {pt}")

    # perimeter_lf: sanity vs footprint
    pl = pq.get('perimeter_lf')
    fp_for_perim = pq.get('footprint_sf')
    if (pl is not None and fp_for_perim is not None
            and isinstance(pl, (int, float)) and isinstance(fp_for_perim, (int, float))
            and fp_for_perim > 0):
        max_perim = 4 * math.sqrt(fp_for_perim) * 2
        if pl > max_perim:
            _pop_bad('perimeter_lf', pl, f"exceeds 2x square perimeter for {fp_for_perim} SF")

    # window_count: vs unit_count
    wc = pq.get('window_count')
    uc_check = pq.get('unit_count') or 1
    if wc is not None and isinstance(wc, (int, float)) and wc > uc_check * 20:
        _pop_bad('window_count', wc, f"exceeds {uc_check} units x 20")

    # floor_to_floor_height_ft: 7–30
    fh = pq.get('floor_to_floor_height_ft')
    if fh is not None and (not isinstance(fh, (int, float)) or fh < 7 or fh > 30):
        _pop_bad('floor_to_floor_height_ft', fh, "must be 7-30 ft")

    if findings:
        pq.setdefault('_validation_findings', []).extend(findings)

    return pq


def process_files(
    file_paths: list[Path],
    api_key: str,
    progress_callback: Optional[Callable[[str], None]] = None,
    gemini_api_key: str = "",
) -> tuple[list[SubmissionRequirement], list[TradeAndScope], dict]:
    """
    Process a list of files (PDFs, images, DOCX, EML, XLSX, CSV, TXT):
    extract requirements and trades from every file.

    Args:
        file_paths: List of Path objects pointing to supported files.
        api_key: Anthropic API key.
        progress_callback: Optional callable that receives progress message strings.

    Returns:
        (requirements, trades, stats) where stats is a dict with token counts,
        costs, page counts, and per-file breakdowns.
    """
    def emit(msg: str):
        if progress_callback:
            progress_callback(msg)

    client = anthropic.Anthropic(api_key=api_key)

    all_requirements: list[SubmissionRequirement] = []
    all_trades: list[TradeAndScope] = []

    file_stats: dict[str, dict] = {}
    global_stats = {
        "total_pages": 0,
        "input_tokens": 0,
        "output_tokens": 0,
        "api_calls": 0,
        "failed_pages": [],
    }

    image_exts = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}
    text_exts = {".docx", ".doc", ".xlsx", ".xls", ".csv", ".eml", ".txt", ".rtf"}

    for file_path in file_paths:
        if not file_path.exists():
            emit(f"[WARNING] File not found, skipping: {file_path}")
            continue

        ext = file_path.suffix.lower()
        emit(f"Opening {file_path.name}...")

        if ext == ".pdf":
            reqs, trades = _process_pdf(client, file_path, emit, global_stats, file_stats, gemini_api_key=gemini_api_key)
        elif ext in image_exts:
            reqs, trades = _process_image(client, file_path, emit, global_stats, file_stats, gemini_api_key=gemini_api_key)
        elif ext in text_exts:
            reqs, trades = _process_text_file(client, file_path, emit, global_stats, file_stats)
        else:
            emit(f"[WARNING] Unsupported file type: {ext}, skipping {file_path.name}")
            continue

        all_requirements.extend(reqs)
        all_trades.extend(trades)

    emit(f"[SCALAR INPUT] {len(all_trades)} trades feeding scalar extraction")

    # Dedup
    emit("Deduplicating results...")
    all_requirements = dedup_requirements(all_requirements)
    all_trades = dedup_trades(all_trades)

    # Tag addenda and detect conflicts
    emit("Scanning for addenda...")
    all_trades = tag_addenda(all_trades)
    addenda_findings = detect_addenda_conflicts(all_trades, emit)

    # Compute costs
    input_cost = (global_stats["input_tokens"] / 1_000_000) * INPUT_RATE
    output_cost = (global_stats["output_tokens"] / 1_000_000) * OUTPUT_RATE
    total_cost = input_cost + output_cost

    consolidated = consolidate_trades(all_trades)

    stats = {
        **global_stats,
        "input_cost": input_cost,
        "output_cost": output_cost,
        "total_cost": total_cost,
        "file_stats": file_stats,
        "num_requirements": len(all_requirements),
        "num_trades": len(consolidated),
        "num_raw_trade_items": len(all_trades),
        "addenda_findings": addenda_findings,
    }

    emit(f"Done! {len(all_requirements)} requirements, {len(consolidated)} trades extracted.")

    return all_requirements, all_trades, stats


# Backwards compatibility alias
process_pdfs = process_files


# ─── Excel builder ───────────────────────────────────────────────────────────


def build_excel_bytes(
    requirements: list[SubmissionRequirement],
    trades: list[TradeAndScope],
    sov_data: Optional[dict] = None,
    failed_pages: Optional[list[str]] = None,
    project_quantities: Optional[dict] = None,
    addenda_findings: Optional[list[dict]] = None,
    progress_callback: Optional[Callable[[str], None]] = None,
) -> bytes:
    """Build the Excel workbook in memory and return it as bytes.

    Produces:
      Tab 1 — Buyout Summary    (buyout-style with financials)
      Tab 2 — Scope Details     (consolidated scope per trade)
      Tab 3 — Submission Reqs   (permits, submittals, inspections)
      Tab 4 — Source Trace      (raw per-page evidence)
    """
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side, numbers
    from openpyxl.utils import get_column_letter

    consolidated = consolidate_trades(trades)

    # ── Project type & site detection ─────────────────────────────────────
    # Build a lightweight extraction dict from the raw trade data for detection
    _raw_text_parts = []
    for t in trades:
        _raw_text_parts.append(t.trade)
        _raw_text_parts.append(t.scope_description)
        if t.vendor_name:
            _raw_text_parts.append(t.vendor_name)
        _raw_text_parts.append(t.evidence)
    for r in requirements:
        _raw_text_parts.append(r.category)
        _raw_text_parts.append(r.requirement)
    _extraction_for_detection = {"raw_text": " ".join(_raw_text_parts)}
    project_type_result = detect_project_type(_extraction_for_detection)

    site_result = _detect_site_scope(consolidated, trades)

    # ── Rate table profile matching ──────────────────────────────────────
    rate_profile = _match_rate_profile(_RATE_TABLES, project_quantities)

    def emit(msg: str):
        if progress_callback:
            progress_callback(msg)

    wb = openpyxl.Workbook()

    # ── Styles ────────────────────────────────────────────────────────────
    navy_fill = PatternFill(start_color="1A2332", end_color="1A2332", fill_type="solid")
    orange_fill = PatternFill(start_color="E8722A", end_color="E8722A", fill_type="solid")
    light_gray_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    summary_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    amber_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    site_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    blue_rate_font = Font(name="Calibri", bold=True, size=11, color="0000FF")
    light_blue_fill = PatternFill(start_color="DDEEFF", end_color="DDEEFF", fill_type="solid")
    dark_blue_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    dark_green_fill = PatternFill(start_color="375623", end_color="375623", fill_type="solid")
    project_total_fill = PatternFill(start_color="1F2D3D", end_color="1F2D3D", fill_type="solid")

    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    title_font = Font(name="Calibri", bold=True, color="FFFFFF", size=14)
    bold_font = Font(name="Calibri", bold=True, size=11)
    normal_font = Font(name="Calibri", size=11)
    money_fmt = '#,##0'
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    top_align = Alignment(vertical="top")
    wrap_top = Alignment(vertical="top", wrap_text=True)

    # ── Scalar Gate: Show extracted quantities with confidence ────────────
    if project_quantities:
        _conf = project_quantities.get('_confidence', {})
        _conf_emoji = {'high': '\U0001f7e2', 'medium': '\U0001f7e1', 'low': '\u26aa'}

        ws_scalars = wb.active
        ws_scalars.title = "Project Scalars"
        ws_scalars.column_dimensions['A'].width = 25
        ws_scalars.column_dimensions['B'].width = 18
        ws_scalars.column_dimensions['C'].width = 12
        ws_scalars.column_dimensions['D'].width = 18
        ws_scalars.column_dimensions['E'].width = 50

        for col, title in enumerate(["Scalar", "Extracted Value", "Confidence",
                                       "Your Override", "Notes"], 1):
            cell = ws_scalars.cell(row=1, column=col, value=title)
            cell.font = header_font
            cell.fill = navy_fill

        scalar_display = [
            ('unit_count', 'Unit Count', 'units'),
            ('floor_count', 'Floor Count', 'floors'),
            ('total_building_sf', 'Total Building SF', 'SF'),
            ('footprint_sf', 'Footprint SF', 'SF'),
            ('perimeter_lf', 'Perimeter LF', 'LF'),
            ('project_type', 'Project Type', ''),
            ('construction_type', 'Construction Type', ''),
            ('roof_type', 'Roof Type', ''),
            ('hvac_system_type', 'HVAC System Type', ''),
            ('foundation_type', 'Foundation Type', ''),
        ]

        for row_idx, (key, label, unit_hint) in enumerate(scalar_display, 2):
            val = project_quantities.get(key, 'NOT EXTRACTED')
            conf = _conf.get(key, 'low')
            _emoji = _conf_emoji.get(conf, '\u26aa')

            ws_scalars.cell(row=row_idx, column=1, value=label).font = bold_font

            if isinstance(val, (int, float)):
                display_val = f"{val:,.0f} {unit_hint}".strip()
            else:
                display_val = str(val) if val is not None else 'NOT EXTRACTED'
            ws_scalars.cell(row=row_idx, column=2, value=display_val).font = normal_font

            conf_cell = ws_scalars.cell(row=row_idx, column=3, value=f"{_emoji} {conf.upper()}")
            conf_cell.font = normal_font
            if conf == 'low':
                conf_cell.fill = amber_fill
            elif conf == 'high':
                conf_cell.fill = green_fill

            # Column D left blank for estimator override
            ws_scalars.cell(row=row_idx, column=4, value="").font = normal_font

            # Notes column
            note = ""
            if conf == 'low':
                note = "\u26a0\ufe0f REQUIRES CONFIRMATION \u2014 pricing depends on this value"
            elif conf == 'medium':
                note = "Please verify"
            ws_scalars.cell(row=row_idx, column=5, value=note).font = normal_font

        # Warning banner if any scalar is low confidence
        if project_quantities.get('_scalar_gate_required', False):
            banner_row = len(scalar_display) + 3
            banner_cell = ws_scalars.cell(
                row=banner_row, column=1,
                value="\u26a0\ufe0f LOW-CONFIDENCE SCALARS DETECTED \u2014 Review values above before trusting pricing"
            )
            banner_cell.font = Font(name="Calibri", bold=True, size=12, color="CC0000")
            ws_scalars.merge_cells(
                start_row=banner_row, start_column=1,
                end_row=banner_row, end_column=5
            )

        # Step B3: Emit scalar gate status
        emit(f"Scalar Gate: {len([v for v in _conf.values() if v == 'high'])} HIGH / "
             f"{len([v for v in _conf.values() if v == 'medium'])} MEDIUM / "
             f"{len([v for v in _conf.values() if v == 'low'])} LOW confidence")

        if project_quantities.get('_scalar_gate_required', False):
            emit("SCALAR GATE: Low-confidence values detected \u2014 "
                 "estimator review required before trusting pricing")

    # ── TAB 1: Buyout Summary ─────────────────────────────────────────────
    ws = wb.create_sheet("Buyout Summary") if project_quantities else wb.active
    if not project_quantities:
        ws.title = "Buyout Summary"
    NUM_COLS = 7  # A through G

    # Column layout: DIV | TRADE | SCOPE | LOW BIDDER | BUDGET / SOV | VARIANCE | NOTES
    col_headers = [
        ("A", "DIV", 7),
        ("B", "TRADE", 34),
        ("C", "SCOPE", 62),
        ("D", "LOW BIDDER", 22),
        ("E", "BUDGET / SOV", 18),
        ("F", "VARIANCE", 14),
        ("G", "NOTES", 36),
    ]

    # Match SOV values from reference buyout if provided
    sov_matched = {}
    sov_summary = {"ohp": None, "bond": None, "permit": None}
    if sov_data:
        sov_matched = _match_sov_to_trades(
            consolidated, sov_data.get("trade_sov", {})
        )
        sov_summary["ohp"] = sov_data.get("ohp")
        sov_summary["bond"] = sov_data.get("bond")
        sov_summary["permit"] = sov_data.get("permit")

    # Masonry guard: only include if rate profile has masonry quote OR extraction has a dollar budget
    _rt_has_masonry = bool(rate_profile.get('masonry'))

    # Filter to trades with pricing data OR pricing engine coverage
    buyout_rows = [
        row for row in consolidated
        if (row["trade"] != "Masonry"
            or _rt_has_masonry
            or (row.get("budget") is not None and row.get("budget", 0) > 0))
        and (row.get("budget") is not None
             or row["trade"] in sov_matched
             or get_baseline_quantities(row["trade"], _extraction_for_detection,
                                        project_quantities=project_quantities))
    ]
    # Safety fallback: if nothing has pricing, show everything
    if not buyout_rows:
        buyout_rows = consolidated

    # Row 1: title bar
    ws.merge_cells("A1:G1")
    title_cell = ws["A1"]
    title_cell.value = "BuildBrain Buyout Summary"
    title_cell.font = title_font
    title_cell.fill = navy_fill
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 40

    # Row 2: PROJECT TYPE banner (when ADA/specialty detected)
    header_row = 2
    if project_type_result["project_type"] != "STANDARD_RESIDENTIAL":
        banner_text = (
            f"\u26a0\ufe0f PROJECT TYPE: {project_type_result['project_type'].replace('_', ' ')}  |  "
            f"Labor Multiplier: per-trade (see Notes)  |  "
            f"{', '.join(project_type_result['trigger_items'][:3])}"
        )
        ws.merge_cells("A2:G2")
        banner_cell = ws["A2"]
        banner_cell.value = banner_text
        banner_cell.font = Font(name="Calibri", bold=True, size=11)
        banner_cell.fill = amber_fill
        banner_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws.row_dimensions[2].height = 32
        header_row = 3

    # Column headers
    for col_letter, label, width in col_headers:
        cell = ws[f"{col_letter}{header_row}"]
        cell.value = label
        cell.font = header_font
        cell.fill = navy_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
        ws.column_dimensions[col_letter].width = width
    ws.row_dimensions[header_row].height = 30

    # Data rows start after header
    data_start = header_row + 1

    # ── Site Work / Civil row (always first trade row) ────────────────
    site_row_idx = data_start
    site_in_buyout = site_result["has_site_scope"] and site_result["site_trade"] is not None
    # Check if Sitework is already in buyout_rows
    site_already_in_buyout = any(
        r["trade"] in ("Sitework", "Paving & Striping", "Site Concrete")
        for r in buyout_rows
    )
    if not site_already_in_buyout:
        # Check user-provided site budget FIRST
        _user_site = (project_quantities or {}).get('site_work_budget')
        if _user_site and isinstance(_user_site, (int, float)) and _user_site > 0:
            site_override = round(_user_site)
            site_override_note = f"User-provided site budget: ${site_override:,.0f}"
        else:
            # Fall back to TRADE_MAP / rate_table logic
            site_items = get_baseline_quantities("SITE WORK / CIVIL", _extraction_for_detection,
                                                 project_quantities=project_quantities)
            site_override = next((item.get("override_amount") for item in site_items
                                 if item.get("override_amount")), 0)
            site_override_note = next((item.get("note") for item in site_items
                                      if item.get("note") and item.get("override_amount")),
                                     site_result["flag_note"])
        ws.cell(row=site_row_idx, column=1, value="02000").font = bold_font
        ws.cell(row=site_row_idx, column=2, value="SITE WORK / CIVIL").font = bold_font
        ws.cell(row=site_row_idx, column=3, value="See Scope Details tab").alignment = wrap_top
        ws.cell(row=site_row_idx, column=5, value=site_override).number_format = money_fmt
        ws.cell(row=site_row_idx, column=7, value=site_override_note).alignment = wrap_top
        ws.cell(row=site_row_idx, column=7).font = Font(name="Calibri", size=9, italic=True)
        for col in range(1, NUM_COLS + 1):
            ws.cell(row=site_row_idx, column=col).fill = site_fill
            ws.cell(row=site_row_idx, column=col).border = thin_border
        ws.row_dimensions[site_row_idx].height = 36
        data_start = site_row_idx + 1
    for i, row_data in enumerate(buyout_rows):
        r = data_start + i
        ws.cell(row=r, column=1, value=int(row_data["csi_division"])).font = bold_font
        ws.cell(row=r, column=2, value=row_data["trade"]).font = bold_font

        # Brief scope: first 3 items, truncated
        scope_lines = row_data["scope"].split("\n")
        brief = "; ".join(line.split(". ", 1)[-1] for line in scope_lines[:3])
        if len(scope_lines) > 3:
            brief += f" (+{len(scope_lines) - 3} more)"
        ws.cell(row=r, column=3, value=brief).alignment = wrap_top

        # LOW BIDDER (D): vendor name of lowest bid
        low_bidder_cell = ws.cell(row=r, column=4)
        bids = row_data.get("bids", [])
        if bids:
            low_bidder_cell.value = bids[0][0]  # vendor name of lowest bid
        low_bidder_cell.alignment = top_align

        # BUDGET / SOV (E): from reference buyout, extracted budget, or pricing engine
        trade_name = row_data["trade"]
        budget_val = None
        note_val = None

        # General Requirements is calculated AFTER all other trades — defer it
        if trade_name == "General Requirements":
            budget_val = 0  # placeholder, recalculated below
            note_val = "Calculated as % of direct building cost (see below)"
            ws.cell(row=r, column=5, value=0).number_format = money_fmt
        else:
            # Tier 1: SOV match from uploaded reference
            if trade_name in sov_matched:
                budget_val = sov_matched[trade_name]

            # Tier 2: Claude extracted dollar amount from PDFs
            if budget_val is None and row_data.get("budget") is not None:
                extracted_budget = row_data["budget"]
                if isinstance(extracted_budget, (int, float)) and extracted_budget > 500000:
                    # Likely a project total or contract amount — do not use
                    budget_val = None
                    note_val = (
                        f"\u26a0\ufe0f LARGE DOLLAR AMOUNT FOUND IN DOCUMENTS: "
                        f"${extracted_budget:,.0f} \u2014 likely a contract total "
                        f"or budget figure, NOT a trade sub quote. "
                        f"Do not use as pricing. Get sub quotes."
                    )
                else:
                    budget_val = extracted_budget

            # Tier 3: Calculate baseline using wage rates + material DB
            if budget_val is None or budget_val == 0:
                work_items = get_baseline_quantities(trade_name, _extraction_for_detection,
                                                     project_quantities=project_quantities)
                if work_items:
                    # Check for override_amount (placeholder trades)
                    override = next((item.get("override_amount") for item in work_items
                                    if item.get("override_amount")), None)
                    override_note = next((item.get("note") for item in work_items
                                         if item.get("note") and item.get("override_amount")), None)

                    if override is not None:
                        # Placeholder trade — write directly, skip calculate_trade_estimate
                        budget_val = override
                        note_val = override_note

                        # Rate table override — replace TRADE_MAP hardcoded rate
                        _rt_bv, _rt_nv = _apply_rate_table_override(
                            trade_name, rate_profile, project_quantities)
                        if _rt_bv is not None:
                            budget_val = _rt_bv
                            note_val = _rt_nv

                    else:
                        # Normal trade — run the pricing engine
                        wage_regime = "CT_DOL_RESIDENTIAL"
                        if isinstance(project_type_result, dict) and "regime" in project_type_result:
                            wage_regime = project_type_result["regime"]

                        trade_estimate = calculate_trade_estimate(
                            trade_name=trade_name,
                            work_items=work_items,
                            wage_regime=wage_regime
                        )

                        budget_val = trade_estimate["total_material"] + trade_estimate["total_labor"]
                        note_val = (
                            f"Calculated baseline | Material: ${trade_estimate['total_material']:,.0f} "
                            f"+ Labor: ${trade_estimate['total_labor']:,.0f} "
                            f"= ${budget_val:,.0f} | \u00b115% accuracy | Sub quote overrides this"
                        )

                        # Rate table override for normal trades
                        _rt_bv, _rt_nv = _apply_rate_table_override(
                            trade_name, rate_profile, project_quantities)
                        if _rt_bv is not None:
                            budget_val = _rt_bv
                            note_val = _rt_nv

                else:
                    budget_val = 0
                    note_val = "\u26a0\ufe0f No pricing found \u2014 estimator must price this trade"
            elif note_val is None:
                note_val = "Extracted from documents or SOV match"

            # Sitework floor safeguard — $102 extraction artifact
            if trade_name in ("Sitework", "SITE WORK / CIVIL",
                              "Paving & Striping", "Landscaping"):
                _pq = project_quantities or {}
                _tsf = _pq.get('total_building_sf') or 1000
                _min_site = _tsf * 5  # absolute minimum $5/SF
                if budget_val is not None and budget_val < _min_site:
                    budget_val = None  # force to Tier 3 placeholder

            # If sitework has no pricing after all tiers, use placeholder
            if trade_name in ("Sitework", "SITE WORK / CIVIL") and \
               (budget_val is None or budget_val == 0):
                _pq = project_quantities or {}
                _tsf = _pq.get('total_building_sf') or 1000
                _pt = _pq.get('project_type', 'unknown')
                if _pt in ('multi_family', 'mixed_use', 'commercial'):
                    _site_rate = 25.00  # urban infill
                else:
                    _site_rate = 15.00  # suburban residential
                budget_val = round(_tsf * _site_rate)
                note_val = (
                    f"\u26a0\ufe0f SITEWORK PLACEHOLDER: ${_site_rate}/SF \u00d7 "
                    f"{_tsf:,.0f} SF = ${budget_val:,.0f} | "
                    f"Includes grading, utilities, paving, landscaping. "
                    f"Urban infill range: $15-35/SF. "
                    f"GET CIVIL SUB QUOTE before bid."
                )

            # Low-end floor safeguard — flag suspiciously low extracted costs
            _MIN_COST_PER_SF = {
                'building concrete': 3.0, 'site concrete': 3.0,
                'structural steel': 4.0,
                'hvac': 5.0,
                'electrical': 4.0,
                'plumbing': 2.5,
                'roofing': 2.0,
                'insulation': 1.5,
                'drywall': 2.0,
                'flooring': 1.5,
                'painting': 0.75,
            }
            _pq_for_floor = project_quantities or {}
            _total_sf_for_floor = _pq_for_floor.get('total_building_sf') or 0
            if (budget_val and isinstance(budget_val, (int, float))
                    and budget_val > 0 and _total_sf_for_floor > 0):
                _trade_lower = trade_name.lower()
                _min_rate = next(
                    (v for k, v in _MIN_COST_PER_SF.items() if k in _trade_lower),
                    1.0
                )
                _floor_min = _min_rate * _total_sf_for_floor * 0.10
                if budget_val < _floor_min:
                    _floor_warning = (
                        f"\u26a0\ufe0f LOW COST FLAG: Extracted ${budget_val:,.0f} "
                        f"appears low for {_total_sf_for_floor:,.0f} SF project. "
                        f"Expected minimum ~${_floor_min:,.0f} "
                        f"({_min_rate}/SF x 10%). Using extracted value \u2014 verify."
                    )
                    note_val = _floor_warning + (" | " + note_val if note_val else "")

            # ── Apply complexity multiplier (renovation/gut rehab/addition) ──
            _complexity_mult = (project_quantities or {}).get('complexity_multiplier', 1.0)
            if (budget_val and isinstance(budget_val, (int, float))
                    and budget_val > 0 and _complexity_mult != 1.0):
                _pre_mult = budget_val
                budget_val = round(budget_val * _complexity_mult)
                _mult_label = (project_quantities or {}).get('construction_type', 'unknown')
                _mult_note = (
                    f"Complexity adj: ${_pre_mult:,.0f} x {_complexity_mult:.2f} "
                    f"({_mult_label}) = ${budget_val:,.0f}"
                )
                note_val = _mult_note + (" | " + note_val if note_val else "")

            # Write to column E (always, never blank)
            ws.cell(row=r, column=5, value=round(budget_val) if budget_val else 0).number_format = money_fmt

        # VARIANCE (F): blank for user
        ws.cell(row=r, column=6, value="").font = normal_font
        # NOTES (G): pricing source info
        existing_g = ws.cell(row=r, column=7).value
        if not existing_g:
            ws.cell(row=r, column=7, value=note_val).font = Font(
                name="Calibri", size=9, italic=True)

        # Alternate row shading
        if i % 2 == 1:
            for col in range(1, NUM_COLS + 1):
                ws.cell(row=r, column=col).fill = light_gray_fill

        # All cells get border
        for col in range(1, NUM_COLS + 1):
            ws.cell(row=r, column=col).border = thin_border

        ws.row_dimensions[r].height = max(30, min(len(scope_lines) * 18, 90))

    data_end = data_start + len(buyout_rows) - 1

    # ── Recalculate General Requirements as % of other direct costs ────
    _pq = project_quantities or {}
    _gr_unit_count = _pq.get('unit_count') or 1
    if _gr_unit_count >= 50:
        _gr_rate = 0.05
    elif _gr_unit_count >= 10:
        _gr_rate = 0.06
    else:
        _gr_rate = 0.08

    # Sum all non-GR trade values written so far
    _other_trades_total = 0
    _gr_row = None
    for _ri in range(data_start, data_end + 1):
        _trade_cell = ws.cell(row=_ri, column=2).value
        _val = ws.cell(row=_ri, column=5).value
        if _trade_cell == "General Requirements":
            _gr_row = _ri
        elif isinstance(_val, (int, float)):
            _other_trades_total += _val

    if _gr_row is not None:
        _gr_budget = round(_other_trades_total * _gr_rate)
        ws.cell(row=_gr_row, column=5, value=_gr_budget).number_format = money_fmt
        _gr_note = (
            f"General Requirements: {_gr_rate:.0%} of "
            f"${_other_trades_total:,.0f} direct cost = "
            f"${_gr_budget:,.0f} | "
            f"Adjust based on project duration and GC overhead structure"
        )
        ws.cell(row=_gr_row, column=7, value=_gr_note).font = Font(
            name="Calibri", size=9, italic=True)

    # ── Python-calculated summary values ─────────────────────────────────
    # Collect all building trade E values (excludes site row)
    building_vals = []
    for r in range(data_start, data_end + 1):
        # Skip if this row is a site trade (CSI div 02)
        _csi_cell = ws.cell(row=r, column=1).value
        if str(_csi_cell or "").startswith("02"):
            continue
        v = ws.cell(row=r, column=5).value
        if isinstance(v, (int, float)):
            building_vals.append(v)
    building_subtotal = sum(building_vals)

    # Site subtotal: read from dedicated site row, or find div-02 trades in data range
    if not site_already_in_buyout:
        site_subtotal = ws.cell(row=site_row_idx, column=5).value or 0
        if not isinstance(site_subtotal, (int, float)):
            site_subtotal = 0
    else:
        # Site trade is within the buyout_rows range — sum div-02 cells
        _site_vals = []
        for r in range(data_start, data_end + 1):
            _csi_cell = ws.cell(row=r, column=1).value
            if str(_csi_cell or "").startswith("02"):
                v = ws.cell(row=r, column=5).value
                if isinstance(v, (int, float)):
                    _site_vals.append(v)
        site_subtotal = sum(_site_vals)

    # GC markup: user-confirmed (pq) > rate_profile > project-type defaults
    _pq_gc = project_quantities or {}
    _gc = rate_profile.get('gc_markups', {})
    _pt = _pq_gc.get('project_type', 'unknown')
    _type_gc_defaults = {
        'single_family':  (0.05, 0.06, 0.05, 0.025),
        'multi_family':   (0.08, 0.10, 0.08, 0.025),
        'mixed_use':      (0.08, 0.10, 0.08, 0.025),
        'commercial':     (0.10, 0.12, 0.10, 0.025),
    }
    _td = _type_gc_defaults.get(_pt, (0.10, 0.12, 0.10, 0.025))
    gc_oh_rate       = _pq_gc.get('gc_conditions_pct') or _gc.get('conditions_pct') or _td[0]
    gc_profit_rate   = _pq_gc.get('gc_profit_pct') or _gc.get('profit_pct') or _td[1]
    contingency_rate = _pq_gc.get('gc_contingency_pct') or _gc.get('contingency_pct') or _td[2]
    permits_rate     = _pq_gc.get('gc_permits_pct') or _gc.get('permits_bond_pct') or _td[3]
    site_fee_rate    = 0.05

    gc_oh          = round(building_subtotal * gc_oh_rate)
    gc_profit      = round(building_subtotal * gc_profit_rate)
    contingency    = round(building_subtotal * contingency_rate)
    permits        = round(building_subtotal * permits_rate)
    building_total = building_subtotal + gc_oh + gc_profit + contingency + permits

    site_fee       = round(site_subtotal * site_fee_rate)
    site_total     = site_subtotal + site_fee
    project_total  = building_total + site_total

    print(f"[EXCEL TOTALS] building_sub={building_subtotal:,.0f} site_sub={site_subtotal:,.0f} "
          f"GC={gc_oh+gc_profit+contingency+permits:,.0f} ({gc_oh_rate+gc_profit_rate+contingency_rate+permits_rate:.1%}) "
          f"project_total={project_total:,.0f}", flush=True)

    currency_fmt = '$#,##0'

    # ── Summary rows below data ──────────────────────────────────────────
    gap = data_end + 2

    # ── SECTION A: BUILDING TRADES ────────────────────────────────────
    bldg_sub_row = gap
    ws.cell(row=bldg_sub_row, column=2, value="BUILDING SUBTOTAL").font = bold_font
    ws.cell(row=bldg_sub_row, column=5, value=building_subtotal).number_format = currency_fmt
    ws.cell(row=bldg_sub_row, column=5).font = bold_font
    ws.cell(row=bldg_sub_row, column=7,
            value=f"Sum of {len(building_vals)} building trades").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=bldg_sub_row, column=col).fill = light_blue_fill
        ws.cell(row=bldg_sub_row, column=col).border = thin_border
    ws.row_dimensions[bldg_sub_row].height = 28

    # GC General Conditions (10%)
    gc_cond_row = bldg_sub_row + 1
    ws.cell(row=gc_cond_row, column=2, value="GC General Conditions").font = bold_font
    ws.cell(row=gc_cond_row, column=5, value=gc_oh).number_format = currency_fmt
    ws.cell(row=gc_cond_row, column=7,
            value=f"${building_subtotal:,.0f} \u00d7 {gc_oh_rate:.0%}").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=gc_cond_row, column=col).border = thin_border
    ws.row_dimensions[gc_cond_row].height = 28

    # GC Profit (12%, green background)
    gc_profit_row = gc_cond_row + 1
    ws.cell(row=gc_profit_row, column=2, value="GC Profit").font = bold_font
    ws.cell(row=gc_profit_row, column=5, value=gc_profit).number_format = currency_fmt
    ws.cell(row=gc_profit_row, column=7,
            value=f"${building_subtotal:,.0f} \u00d7 {gc_profit_rate:.0%}").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=gc_profit_row, column=col).fill = green_fill
        ws.cell(row=gc_profit_row, column=col).border = thin_border
    ws.row_dimensions[gc_profit_row].height = 28

    # Contingency (10%, amber background)
    contingency_row = gc_profit_row + 1
    ws.cell(row=contingency_row, column=2, value="Contingency").font = bold_font
    ws.cell(row=contingency_row, column=5, value=contingency).number_format = currency_fmt
    ws.cell(row=contingency_row, column=7,
            value=f"${building_subtotal:,.0f} \u00d7 {contingency_rate:.0%}").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=contingency_row, column=col).fill = amber_fill
        ws.cell(row=contingency_row, column=col).border = thin_border
    ws.row_dimensions[contingency_row].height = 28

    # Permits + Bond (2.5%)
    permits_bond_row = contingency_row + 1
    ws.cell(row=permits_bond_row, column=2, value="Permits + Bond").font = bold_font
    ws.cell(row=permits_bond_row, column=5, value=permits).number_format = currency_fmt
    ws.cell(row=permits_bond_row, column=7,
            value=f"${building_subtotal:,.0f} \u00d7 {permits_rate:.1%}").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=permits_bond_row, column=col).border = thin_border
    ws.row_dimensions[permits_bond_row].height = 28

    # BUILDING TOTAL row (dark blue background, white bold)
    bldg_total_row = permits_bond_row + 1
    ws.cell(row=bldg_total_row, column=2, value="BUILDING TOTAL").font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    ws.cell(row=bldg_total_row, column=5, value=building_total).number_format = currency_fmt
    ws.cell(row=bldg_total_row, column=5).font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    ws.cell(row=bldg_total_row, column=7,
            value=f"${building_subtotal:,.0f} + markups").font = Font(
        name="Calibri", size=9, italic=True, color="FFFFFF")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=bldg_total_row, column=col).fill = dark_blue_fill
        ws.cell(row=bldg_total_row, column=col).font = Font(
            name="Calibri", bold=True, size=12, color="FFFFFF")
        ws.cell(row=bldg_total_row, column=col).border = thin_border
    ws.row_dimensions[bldg_total_row].height = 32

    # Spacer between sections
    spacer1_row = bldg_total_row + 1
    ws.row_dimensions[spacer1_row].height = 10

    # ── SECTION B: SITE WORK ─────────────────────────────────────────
    site_sub_row = spacer1_row + 1
    ws.cell(row=site_sub_row, column=2, value="SITE WORK SUBTOTAL").font = bold_font
    ws.cell(row=site_sub_row, column=5, value=site_subtotal).number_format = currency_fmt
    ws.cell(row=site_sub_row, column=5).font = bold_font
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=site_sub_row, column=col).fill = site_fill
        ws.cell(row=site_sub_row, column=col).border = thin_border
    ws.row_dimensions[site_sub_row].height = 28

    # GC Management Fee (5%)
    site_fee_row = site_sub_row + 1
    ws.cell(row=site_fee_row, column=2, value="GC Management Fee").font = bold_font
    ws.cell(row=site_fee_row, column=5, value=site_fee).number_format = currency_fmt
    ws.cell(row=site_fee_row, column=7,
            value=f"${site_subtotal:,.0f} \u00d7 {site_fee_rate:.0%} GC mgmt fee").font = Font(
        name="Calibri", size=9, italic=True, color="666666")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=site_fee_row, column=col).border = thin_border
    ws.row_dimensions[site_fee_row].height = 28

    # SITE TOTAL row (dark green background, white bold)
    site_total_row = site_fee_row + 1
    ws.cell(row=site_total_row, column=2, value="SITE TOTAL").font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    ws.cell(row=site_total_row, column=5, value=site_total).number_format = currency_fmt
    ws.cell(row=site_total_row, column=5).font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=site_total_row, column=col).fill = dark_green_fill
        ws.cell(row=site_total_row, column=col).font = Font(
            name="Calibri", bold=True, size=12, color="FFFFFF")
        ws.cell(row=site_total_row, column=col).border = thin_border
    ws.row_dimensions[site_total_row].height = 32

    # Explanatory note below SITE TOTAL
    note_row = site_total_row + 1
    ws.merge_cells(f"C{note_row}:G{note_row}")
    note_cell = ws.cell(row=note_row, column=3)
    note_cell.value = (
        "Civil sub quotes are lump sum and include sub\u2019s overhead & profit. "
        "GC management fee only (5%) applied to site scope. "
        "Building trades include full GC overhead, profit, and contingency."
    )
    note_cell.font = Font(name="Calibri", size=9, italic=True, color="666666")
    note_cell.alignment = wrap_top
    ws.row_dimensions[note_row].height = 36

    # Spacer before project total
    spacer2_row = note_row + 1
    ws.row_dimensions[spacer2_row].height = 10

    # ── SECTION C: PROJECT TOTAL ─────────────────────────────────────
    project_total_row = spacer2_row + 1
    ws.cell(row=project_total_row, column=2, value="PROJECT TOTAL").font = Font(
        name="Calibri", bold=True, size=14, color="FFFFFF")
    ws.cell(row=project_total_row, column=5, value=project_total).number_format = currency_fmt
    ws.cell(row=project_total_row, column=5).font = Font(
        name="Calibri", bold=True, size=14, color="FFFFFF")
    _pt_note = f"Building ${building_total:,.0f} + Site ${site_total:,.0f}"
    _pt_complexity_warn = (project_quantities or {}).get('complexity_warning')
    if _pt_complexity_warn:
        _pt_note += f" | {_pt_complexity_warn}"
    ws.cell(row=project_total_row, column=7, value=_pt_note).font = Font(
        name="Calibri", size=9, italic=True, color="FFFFFF")
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=project_total_row, column=col).fill = project_total_fill
        ws.cell(row=project_total_row, column=col).font = Font(
            name="Calibri", bold=True, size=14, color="FFFFFF")
        ws.cell(row=project_total_row, column=col).border = thin_border
    ws.row_dimensions[project_total_row].height = 36

    # ── Estimate band: LOW / HIGH ────────────────────────────────────
    low_green_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    low_red_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
    variance_note = "Reflects \u00b115% estimating variance for drawing-only takeoff"

    low_row = project_total_row + 1
    low_val = round(project_total * 0.85)
    ws.cell(row=low_row, column=2, value="LOW ESTIMATE (\u221215%)").font = bold_font
    ws.cell(row=low_row, column=5, value=low_val).number_format = currency_fmt
    ws.cell(row=low_row, column=5).font = bold_font
    ws.cell(row=low_row, column=7, value=variance_note).font = Font(
        name="Calibri", size=9, italic=True)
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=low_row, column=col).fill = low_green_fill
        ws.cell(row=low_row, column=col).border = thin_border
    ws.row_dimensions[low_row].height = 28

    high_row = low_row + 1
    high_val = round(project_total * 1.15)
    ws.cell(row=high_row, column=2, value="HIGH ESTIMATE (+15%)").font = bold_font
    ws.cell(row=high_row, column=5, value=high_val).number_format = currency_fmt
    ws.cell(row=high_row, column=5).font = bold_font
    ws.cell(row=high_row, column=7, value=variance_note).font = Font(
        name="Calibri", size=9, italic=True)
    for col in range(1, NUM_COLS + 1):
        ws.cell(row=high_row, column=col).fill = low_red_fill
        ws.cell(row=high_row, column=col).border = thin_border
    ws.row_dimensions[high_row].height = 28

    # Disclaimer note
    disc_row = high_row + 2
    ws.merge_cells(f"B{disc_row}:G{disc_row}")
    disc_cell = ws.cell(row=disc_row, column=2)
    disc_cell.value = (
        "BuildBrain baseline calculated from CT DOL Prevailing Wage rates and "
        "published material costs. Actual sub quotes will vary by market conditions, "
        "project timing, and site-specific factors. This estimate is for budgeting "
        "purposes only \u2014 obtain competitive sub bids before final pricing."
    )
    disc_cell.font = Font(name="Calibri", size=9, italic=True, color="888888")
    disc_cell.alignment = wrap_top
    ws.row_dimensions[disc_row].height = 40

    # Freeze panes (adjust for banner row)
    ws.freeze_panes = f"C{header_row + 1}"
    ws.sheet_properties.tabColor = "1A2332"

    # ── TAB 2: Scope Details ──────────────────────────────────────────────
    ws2 = wb.create_sheet("Scope Details")
    ws2.sheet_properties.tabColor = "E8722A"

    scope_headers = ["DIV", "TRADE", "SCOPE", "SOURCE PAGES"]
    for ci, h in enumerate(scope_headers, 1):
        cell = ws2.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = navy_fill
        cell.border = thin_border

    ws2.column_dimensions["A"].width = 7
    ws2.column_dimensions["B"].width = 34
    ws2.column_dimensions["C"].width = 110
    ws2.column_dimensions["D"].width = 55
    ws2.row_dimensions[1].height = 28

    for i, row_data in enumerate(consolidated):
        r = i + 2
        ws2.cell(row=r, column=1, value=int(row_data["csi_division"])).alignment = top_align
        ws2.cell(row=r, column=2, value=row_data["trade"]).alignment = top_align
        ws2.cell(row=r, column=2).font = bold_font
        scope_cell = ws2.cell(row=r, column=3, value=row_data["scope"])
        scope_cell.alignment = wrap_top
        ws2.cell(row=r, column=4, value=row_data["source_pages"]).alignment = wrap_top
        line_count = row_data["scope"].count("\n") + 1
        ws2.row_dimensions[r].height = max(30, line_count * 18)
        for ci in range(1, 5):
            ws2.cell(row=r, column=ci).border = thin_border

    # ── TAB 3: Submission Requirements ────────────────────────────────────
    ws3 = wb.create_sheet("Submission Requirements")
    ws3.sheet_properties.tabColor = "27AE60"

    req_headers = ["CATEGORY", "REQUIREMENT", "MANDATORY", "SOURCE PDF", "PAGE", "EVIDENCE"]
    req_widths = [22, 70, 14, 36, 8, 70]
    for ci, (h, w) in enumerate(zip(req_headers, req_widths), 1):
        cell = ws3.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = navy_fill
        cell.border = thin_border
        ws3.column_dimensions[get_column_letter(ci)].width = w
    ws3.row_dimensions[1].height = 28

    sorted_reqs = sorted(requirements, key=lambda r: (r.category, r.source_pdf, r.source_page))
    for i, req in enumerate(sorted_reqs):
        r = i + 2
        ws3.cell(row=r, column=1, value=req.category).border = thin_border
        ws3.cell(row=r, column=1).alignment = top_align
        ws3.cell(row=r, column=2, value=req.requirement).border = thin_border
        ws3.cell(row=r, column=2).alignment = wrap_top
        ws3.cell(row=r, column=3, value="Yes" if req.mandatory else "No").border = thin_border
        ws3.cell(row=r, column=3).alignment = Alignment(horizontal="center", vertical="top")
        ws3.cell(row=r, column=4, value=req.source_pdf).border = thin_border
        ws3.cell(row=r, column=4).alignment = wrap_top
        ws3.cell(row=r, column=5, value=req.source_page).border = thin_border
        ws3.cell(row=r, column=5).alignment = Alignment(horizontal="center", vertical="top")
        ws3.cell(row=r, column=6, value=req.evidence).border = thin_border
        ws3.cell(row=r, column=6).alignment = wrap_top
        # Row height based on longest text in requirement or evidence
        max_len = max(len(req.requirement), len(req.evidence))
        ws3.row_dimensions[r].height = max(28, min(int(max_len / 60) * 18 + 28, 80))

    # ── TAB 4: Source Trace ───────────────────────────────────────────────
    ws4 = wb.create_sheet("Source Trace")
    ws4.sheet_properties.tabColor = "7F8C9B"

    trace_headers = ["TYPE", "CATEGORY / TRADE", "DETAIL", "SOURCE PDF", "PAGE", "EVIDENCE"]
    trace_widths = [16, 34, 70, 36, 8, 70]
    for ci, (h, w) in enumerate(zip(trace_headers, trace_widths), 1):
        cell = ws4.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = navy_fill
        cell.border = thin_border
        ws4.column_dimensions[get_column_letter(ci)].width = w
    ws4.row_dimensions[1].height = 28

    all_traces = []
    for req in requirements:
        all_traces.append(("Requirement", req.category, req.requirement,
                           req.source_pdf, req.source_page, req.evidence))
    for t in trades:
        all_traces.append(("Trade", t.trade, t.scope_description,
                           t.source_pdf, t.source_page, t.evidence))
    all_traces.sort(key=lambda x: (x[3], x[4], x[0]))

    for i, (rtype, cat, detail, pdf, page, ev) in enumerate(all_traces):
        r = i + 2
        ws4.cell(row=r, column=1, value=rtype).border = thin_border
        ws4.cell(row=r, column=1).alignment = top_align
        ws4.cell(row=r, column=2, value=cat).border = thin_border
        ws4.cell(row=r, column=2).alignment = top_align
        ws4.cell(row=r, column=3, value=detail).border = thin_border
        ws4.cell(row=r, column=3).alignment = wrap_top
        ws4.cell(row=r, column=4, value=pdf).border = thin_border
        ws4.cell(row=r, column=4).alignment = wrap_top
        ws4.cell(row=r, column=5, value=page).border = thin_border
        ws4.cell(row=r, column=5).alignment = Alignment(horizontal="center", vertical="top")
        ws4.cell(row=r, column=6, value=ev).border = thin_border
        ws4.cell(row=r, column=6).alignment = wrap_top
        max_len = max(len(detail), len(ev))
        ws4.row_dimensions[r].height = max(26, min(int(max_len / 55) * 18 + 26, 70))

    # ── TAB 5: Calculation Detail ───────────────────────────────────────
    ws5 = wb.create_sheet("Calculation Detail")
    ws5.sheet_properties.tabColor = "2E75B6"

    calc_headers = [
        ("A", "TRADE", 22), ("B", "LINE ITEM DESCRIPTION", 35),
        ("C", "QUANTITY", 10), ("D", "UNIT", 8),
        ("E", "UNIT COST (MAT)", 14), ("F", "MATERIAL TOTAL", 14),
        ("G", "LABOR HRS/UNIT", 14), ("H", "WAGE RATE ($/HR)", 14),
        ("I", "LABOR TOTAL", 14), ("J", "LINE TOTAL", 14),
        ("K", "SOURCE", 25),
    ]
    for col_letter, label, width in calc_headers:
        cell = ws5[f"{col_letter}1"]
        cell.value = label
        cell.font = header_font
        cell.fill = navy_fill
        cell.border = thin_border
        ws5.column_dimensions[col_letter].width = width
    ws5.row_dimensions[1].height = 30

    # Row 2: explanatory note
    ws5.merge_cells("A2:K2")
    note = ws5["A2"]
    note.value = (
        "All quantities derived from uploaded drawings. "
        "Wage rates: CT DOL Prevailing Wage ID# 25-6853. "
        "Material prices: BuildBrain material database. "
        "Labor hours: RS Means residential productivity rates. "
        "This sheet is auto-generated \u2014 sub quotes override calculated baselines."
    )
    note.font = Font(name="Calibri", size=9, italic=True, color="666666")
    note.alignment = wrap_top
    ws5.row_dimensions[2].height = 32

    calc_row = 3
    grand_material = 0
    grand_labor = 0

    # Iterate trades that have baseline quantities
    for row_data in buyout_rows:
        t_name = row_data["trade"]
        work_items = get_baseline_quantities(t_name, _extraction_for_detection,
                                             project_quantities=project_quantities)
        if not work_items:
            continue

        # Check for override-only trades (skip detail for placeholders)
        has_real_items = any(
            item.get("work_item") or item.get("material_key")
            for item in work_items
            if not item.get("override_amount")
        )
        if not has_real_items:
            # Placeholder trade — write one summary row
            override_amt = next(
                (item.get("override_amount") for item in work_items if item.get("override_amount")), 0
            )
            ws5.cell(row=calc_row, column=1, value=t_name).font = bold_font
            ws5.cell(row=calc_row, column=2, value=work_items[0].get("description", "")).font = normal_font
            ws5.cell(row=calc_row, column=10, value=override_amt).number_format = currency_fmt
            ws5.cell(row=calc_row, column=11, value="Placeholder \u2014 get sub quote").font = Font(
                name="Calibri", size=9, italic=True)
            for c in range(1, 12):
                ws5.cell(row=calc_row, column=c).border = thin_border
            grand_material += override_amt
            calc_row += 2  # blank row after
            continue

        trade_mat_total = 0
        trade_labor_total = 0
        first_row_of_trade = calc_row

        for item in work_items:
            if item.get("override_amount"):
                continue

            work_item_key = item.get("work_item")
            qty = item.get("quantity", 0)
            mat_key = item.get("material_key")
            mat_qty = item.get("material_quantity", qty)
            desc = item.get("description", "")

            # Material calc
            mat_unit_cost = 0
            mat_total = 0
            mat_unit = ""
            if mat_key and mat_qty > 0:
                mat_data = get_material_price(mat_key)
                if mat_data:
                    mat_unit_cost = mat_data["price"]
                    mat_total = mat_qty * mat_unit_cost
                    mat_unit = mat_data["unit"]
                    trade_mat_total += mat_total

            # Labor calc
            labor_hrs_per_unit = 0
            wage_rate = 0
            labor_total = 0
            labor_unit = ""
            if work_item_key and qty > 0:
                from productivity_rates import PRODUCTIVITY_RATES
                prod = PRODUCTIVITY_RATES.get(work_item_key)
                if prod:
                    labor_hrs_per_unit = prod["hours"]
                    labor_unit = prod["unit"]
                    wage_rate = get_wage(prod["trade"])
                    labor_total = qty * labor_hrs_per_unit * wage_rate
                    trade_labor_total += labor_total

            line_total = mat_total + labor_total
            unit_label = labor_unit or mat_unit or ""

            # Write row
            if calc_row == first_row_of_trade:
                ws5.cell(row=calc_row, column=1, value=t_name).font = bold_font
            ws5.cell(row=calc_row, column=2, value=desc).font = normal_font
            ws5.cell(row=calc_row, column=2).alignment = wrap_top
            if qty > 0:
                ws5.cell(row=calc_row, column=3, value=round(qty, 1)).number_format = '#,##0'
            elif mat_qty > 0:
                ws5.cell(row=calc_row, column=3, value=round(mat_qty, 1)).number_format = '#,##0'
            ws5.cell(row=calc_row, column=4, value=unit_label)
            if mat_unit_cost > 0:
                ws5.cell(row=calc_row, column=5, value=round(mat_unit_cost, 2)).number_format = '$#,##0.00'
            if mat_total > 0:
                ws5.cell(row=calc_row, column=6, value=round(mat_total)).number_format = currency_fmt
            if labor_hrs_per_unit > 0:
                ws5.cell(row=calc_row, column=7, value=labor_hrs_per_unit).number_format = '0.000'
            if wage_rate > 0:
                ws5.cell(row=calc_row, column=8, value=round(wage_rate, 2)).number_format = '$#,##0.00'
            if labor_total > 0:
                ws5.cell(row=calc_row, column=9, value=round(labor_total)).number_format = currency_fmt
            ws5.cell(row=calc_row, column=10, value=round(line_total)).number_format = currency_fmt
            ws5.cell(row=calc_row, column=11, value="material_db + CT DOL wages").font = Font(
                name="Calibri", size=9, italic=True)
            for c in range(1, 12):
                ws5.cell(row=calc_row, column=c).border = thin_border
            calc_row += 1

        # Trade subtotal row
        ws5.cell(row=calc_row, column=2,
                 value=f"TRADE SUBTOTAL \u2014 {t_name}").font = bold_font
        ws5.cell(row=calc_row, column=6, value=round(trade_mat_total)).number_format = currency_fmt
        ws5.cell(row=calc_row, column=6).font = bold_font
        ws5.cell(row=calc_row, column=9, value=round(trade_labor_total)).number_format = currency_fmt
        ws5.cell(row=calc_row, column=9).font = bold_font
        trade_total = round(trade_mat_total + trade_labor_total)
        ws5.cell(row=calc_row, column=10, value=trade_total).number_format = currency_fmt
        ws5.cell(row=calc_row, column=10).font = bold_font
        for c in range(1, 12):
            ws5.cell(row=calc_row, column=c).fill = light_gray_fill
            ws5.cell(row=calc_row, column=c).border = thin_border
        calc_row += 2  # blank row between trades

        grand_material += trade_mat_total
        grand_labor += trade_labor_total

    # ── Grand totals at bottom ──
    grand_direct = round(grand_material + grand_labor)

    ws5.cell(row=calc_row, column=2, value="PROJECT DIRECT COSTS").font = bold_font
    ws5.cell(row=calc_row, column=6, value=round(grand_material)).number_format = currency_fmt
    ws5.cell(row=calc_row, column=6).font = bold_font
    ws5.cell(row=calc_row, column=9, value=round(grand_labor)).number_format = currency_fmt
    ws5.cell(row=calc_row, column=9).font = bold_font
    ws5.cell(row=calc_row, column=10, value=grand_direct).number_format = currency_fmt
    ws5.cell(row=calc_row, column=10).font = bold_font
    for c in range(1, 12):
        ws5.cell(row=calc_row, column=c).fill = light_blue_fill
        ws5.cell(row=calc_row, column=c).border = thin_border
    calc_row += 1

    # GC markups (use rate_profile if matched, else residential defaults)
    markup_items = [
        (f"GC General Conditions ({gc_oh_rate:.0%})", gc_oh_rate),
        (f"GC Profit ({gc_profit_rate:.0%})", gc_profit_rate),
        (f"Contingency ({contingency_rate:.0%})", contingency_rate),
        (f"Permits + Bond ({permits_rate:.1%})", permits_rate),
    ]
    markup_total = 0
    for label, rate in markup_items:
        amt = round(grand_direct * rate)
        markup_total += amt
        ws5.cell(row=calc_row, column=2, value=label).font = normal_font
        ws5.cell(row=calc_row, column=10, value=amt).number_format = currency_fmt
        ws5.cell(row=calc_row, column=11,
                 value=f"${grand_direct:,.0f} \u00d7 {rate:.1%}").font = Font(
            name="Calibri", size=9, italic=True)
        for c in range(1, 12):
            ws5.cell(row=calc_row, column=c).border = thin_border
        calc_row += 1

    # BUILDING TOTAL
    building_calc_total = grand_direct + markup_total
    ws5.cell(row=calc_row, column=2, value="BUILDING TOTAL").font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    ws5.cell(row=calc_row, column=10, value=building_calc_total).number_format = currency_fmt
    ws5.cell(row=calc_row, column=10).font = Font(
        name="Calibri", bold=True, size=12, color="FFFFFF")
    for c in range(1, 12):
        ws5.cell(row=calc_row, column=c).fill = navy_fill
        ws5.cell(row=calc_row, column=c).font = Font(
            name="Calibri", bold=True, size=12, color="FFFFFF")
        ws5.cell(row=calc_row, column=c).border = thin_border

    ws5.freeze_panes = "A3"

    # ── TAB 6: Flags ─────────────────────────────────────────────────────
    _all_flags = []

    # Failed pages
    for page_ref in (failed_pages or []):
        _all_flags.append({
            "severity": "HIGH",
            "flag": "Page could not be parsed",
            "detail": f"{page_ref} failed extraction after 4 attempts. Review manually for missed scope.",
            "source": page_ref,
        })

    # Addenda conflicts
    for finding in (addenda_findings or []):
        _all_flags.append(finding)

    # Validation findings from project quantities
    _pq_val = project_quantities or {}
    for vf in _pq_val.get('_validation_findings', []):
        _all_flags.append({
            "severity": "MEDIUM",
            "flag": "Quantity validation",
            "detail": vf,
            "source": "Project quantities",
        })

    if _all_flags:
        ws6 = wb.create_sheet("Flags")
        ws6.sheet_properties.tabColor = "FF0000"

        flag_headers = ["Severity", "Flag", "Detail", "Source"]
        flag_widths = [12, 30, 70, 30]
        for ci, (h, w) in enumerate(zip(flag_headers, flag_widths), 1):
            cell = ws6.cell(row=1, column=ci, value=h)
            cell.font = header_font
            cell.fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border
            ws6.column_dimensions[get_column_letter(ci)].width = w
        ws6.row_dimensions[1].height = 28

        _severity_colors = {"HIGH": "C00000", "MEDIUM": "E67E22", "LOW": "2980B9"}
        for ri, flag in enumerate(_all_flags, 2):
            sev = flag.get("severity", "MEDIUM")
            sev_color = _severity_colors.get(sev, "666666")
            ws6.cell(row=ri, column=1, value=sev).font = Font(
                name="Calibri", bold=True, size=11, color=sev_color)
            ws6.cell(row=ri, column=1).alignment = Alignment(horizontal="center", vertical="top")
            ws6.cell(row=ri, column=2, value=flag.get("flag", "")).font = bold_font
            ws6.cell(row=ri, column=3, value=flag.get("detail", ""))
            ws6.cell(row=ri, column=3).alignment = wrap_top
            ws6.cell(row=ri, column=4, value=flag.get("source", "")).alignment = top_align
            for c in range(1, 5):
                ws6.cell(row=ri, column=c).border = thin_border

        ws6.freeze_panes = "A2"

    # ── Save ──────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_results_json(
    requirements: list[SubmissionRequirement],
    trades: list[TradeAndScope],
    sov_data: Optional[dict] = None,
    failed_pages: Optional[list[str]] = None,
    project_quantities: Optional[dict] = None,
    addenda_findings: Optional[list[dict]] = None,
) -> dict:
    """Build JSON for the in-browser Procore-style budget viewer.

    Returns dict with: divisions (grouped budget rows), flags, scalars, summary.
    Each division has collapsible rows with editable budget cells.
    """
    consolidated = consolidate_trades(trades)

    _raw_text_parts = []
    for t in trades:
        _raw_text_parts.extend([t.trade, t.scope_description, t.evidence])
        if t.vendor_name:
            _raw_text_parts.append(t.vendor_name)
    for r in requirements:
        _raw_text_parts.extend([r.category, r.requirement])
    _extraction_for_detection = {"raw_text": " ".join(_raw_text_parts)}

    project_type_result = detect_project_type(_extraction_for_detection)
    site_result = _detect_site_scope(consolidated, trades)
    rate_profile = _match_rate_profile(_RATE_TABLES, project_quantities)

    pq = project_quantities or {}

    sov_matched = {}
    if sov_data:
        sov_matched = _match_sov_to_trades(consolidated, sov_data.get("trade_sov", {}))

    # ── Build budget rows ────────────────────────────────────
    all_rows = []

    # Site work row — always add the priced SITE WORK / CIVIL row.
    # Check user-provided site budget FIRST
    _user_site_json = pq.get('site_work_budget')
    _rt_bv = None  # safe default before conditional logic
    _site_source = "placeholder"
    if _user_site_json and isinstance(_user_site_json, (int, float)) and _user_site_json > 0:
        site_override = round(_user_site_json)
        site_note = f"User-provided site budget: ${site_override:,.0f}"
        _site_source = "user"
    else:
        site_items = get_baseline_quantities("SITE WORK / CIVIL", _extraction_for_detection,
                                             project_quantities=pq)
        site_override = next((item.get("override_amount") for item in site_items
                             if item.get("override_amount")), 0)
        site_note = next((item.get("note") for item in site_items
                         if item.get("note") and item.get("override_amount")),
                        site_result.get("flag_note", ""))

        _rt_bv, _rt_nv = _apply_rate_table_override("SITE WORK / CIVIL", rate_profile, pq)
        if _rt_bv is not None:
            site_override = _rt_bv
            site_note = _rt_nv
            _site_source = "rate_table"

    all_rows.append({
        "cost_code": "02-000",
        "div": "02",
        "trade": "SITE WORK / CIVIL",
        "scope_brief": "Grading, utilities, paving, landscaping",
        "low_bidder": None,
        "original_budget": site_override,
        "budget_modifications": 0,
        "revised_budget": site_override,
        "committed_costs": 0,
        "notes": site_note,
        "source": _site_source,
        "scope_items": [],
    })

    # Masonry guard: only include if rate profile has masonry quote OR extraction has a dollar budget
    _rt_has_masonry = bool(rate_profile.get('masonry'))
    print(f"[MASONRY CHECK] rate_profile has masonry={_rt_has_masonry}", flush=True)

    for row_data in consolidated:
        trade_name = row_data["trade"]

        # Guard: skip Masonry unless rate profile prices it or docs had a real dollar amount
        if trade_name == "Masonry":
            _extracted_budget = row_data.get("budget")
            _has_real_budget = _extracted_budget is not None and _extracted_budget > 0
            if not _rt_has_masonry and not _has_real_budget:
                print(f"[MASONRY CHECK] Skipping — no rate profile quote, no extracted budget", flush=True)
                continue
            print(f"[MASONRY CHECK] Including — rt={_rt_has_masonry}, extracted=${_extracted_budget}", flush=True)

        budget_val = None
        note_val = None
        source = "none"

        if trade_name == "General Requirements":
            budget_val = 0
            note_val = "Calculated as % of direct cost"
            source = "calculated"
        else:
            # Tier 1: SOV
            if trade_name in sov_matched:
                budget_val = sov_matched[trade_name]
                note_val = "Matched from reference buyout SOV"
                source = "sov"

            # Tier 2: Extracted
            if budget_val is None and row_data.get("budget") is not None:
                extracted = row_data["budget"]
                if isinstance(extracted, (int, float)) and extracted > 500000:
                    budget_val = None
                    note_val = f"Large amount ${extracted:,.0f} \u2014 likely contract total"
                else:
                    budget_val = extracted
                    note_val = "Extracted from documents"
                    source = "extracted"

            # Tier 3: Rate table or material_db
            if budget_val is None or budget_val == 0:
                work_items = get_baseline_quantities(trade_name, _extraction_for_detection,
                                                     project_quantities=pq)
                if work_items:
                    override = next((item.get("override_amount") for item in work_items
                                    if item.get("override_amount")), None)
                    override_note = next((item.get("note") for item in work_items
                                         if item.get("note") and item.get("override_amount")), None)

                    if override is not None:
                        budget_val = override
                        note_val = override_note
                        source = "placeholder"
                    else:
                        wage_regime = "CT_DOL_RESIDENTIAL"
                        if isinstance(project_type_result, dict) and "regime" in project_type_result:
                            wage_regime = project_type_result["regime"]
                        te = calculate_trade_estimate(
                            trade_name=trade_name, work_items=work_items, wage_regime=wage_regime)
                        budget_val = te["total_material"] + te["total_labor"]
                        note_val = f"Material: ${te['total_material']:,.0f} + Labor: ${te['total_labor']:,.0f}"
                        source = "material_db"

                    _rt_bv, _rt_nv = _apply_rate_table_override(trade_name, rate_profile, pq)
                    if _rt_bv is not None:
                        budget_val = _rt_bv
                        note_val = _rt_nv
                        source = "rate_table"
                else:
                    budget_val = 0
                    note_val = "No pricing \u2014 estimator must price"
                    source = "none"

            # Sitework floor safeguard
            if trade_name in ("Sitework", "SITE WORK / CIVIL", "Paving & Striping", "Landscaping"):
                _tsf = pq.get('total_building_sf') or 1000
                if budget_val is not None and budget_val < _tsf * 5:
                    budget_val = None

            # Complexity multiplier
            _cm = pq.get('complexity_multiplier', 1.0)
            if budget_val and isinstance(budget_val, (int, float)) and budget_val > 0 and _cm != 1.0:
                budget_val = round(budget_val * _cm)
                _label = pq.get('construction_type', 'unknown')
                note_val = f"{_cm:.2f}x ({_label}) | {note_val}" if note_val else f"{_cm:.2f}x ({_label})"

        budget_val = round(budget_val) if budget_val else 0

        scope_lines = row_data["scope"].split("\n")
        brief = "; ".join(line.split(". ", 1)[-1] for line in scope_lines[:2])
        if len(scope_lines) > 2:
            brief += f" (+{len(scope_lines) - 2} more)"

        bids = row_data.get("bids", [])

        all_rows.append({
            "cost_code": f"{row_data['csi_division'][:2]}-{row_data['csi_division']}",
            "div": row_data["csi_division"][:2] if len(str(row_data["csi_division"])) >= 2 else str(row_data["csi_division"]).zfill(2),
            "trade": trade_name,
            "scope_brief": brief,
            "low_bidder": bids[0][0] if bids else None,
            "original_budget": budget_val,
            "budget_modifications": 0,
            "revised_budget": budget_val,
            "committed_costs": 0,
            "notes": note_val or "",
            "source": source,
            "scope_items": [line.split(". ", 1)[-1] if ". " in line else line for line in scope_lines[:10]],
        })

    # Recalculate General Requirements (building trades only, exclude site div 02)
    direct_total = sum(r["original_budget"] for r in all_rows
                       if r["trade"] != "General Requirements" and r["div"] != "02")
    _gr_uc = pq.get('unit_count') or 1
    _gr_rate = 0.05 if _gr_uc >= 50 else (0.06 if _gr_uc >= 10 else 0.08)
    for row in all_rows:
        if row["trade"] == "General Requirements":
            row["original_budget"] = round(direct_total * _gr_rate)
            row["revised_budget"] = row["original_budget"]
            row["notes"] = f"{_gr_rate:.0%} of ${direct_total:,.0f} direct cost"
            break

    # Debug: per-trade values
    for _r in all_rows:
        print(f"[JSON TRADE] div={_r['div']} {_r['trade'][:25]:<25} ${_r['original_budget']:>10,.0f}  src={_r.get('source','?')}", flush=True)

    # ── Group by CSI division ────────────────────────────────
    from collections import OrderedDict
    div_groups = OrderedDict()
    _DIV_NAMES = {
        "01": "General Requirements", "02": "Site Construction",
        "03": "Concrete", "04": "Masonry", "05": "Metals",
        "06": "Wood & Plastics", "07": "Thermal & Moisture",
        "08": "Doors & Windows", "09": "Finishes",
        "10": "Specialties", "11": "Equipment", "12": "Furnishings",
        "13": "Special Construction", "14": "Conveying Systems",
        "15": "Mechanical", "16": "Electrical",
    }

    for row in all_rows:
        div = row["div"]
        if div not in div_groups:
            div_groups[div] = {
                "div": div,
                "name": _DIV_NAMES.get(div, f"Division {div}"),
                "rows": [],
                "subtotal": 0,
            }
        div_groups[div]["rows"].append(row)
        div_groups[div]["subtotal"] += row["original_budget"]

    divisions = list(div_groups.values())

    # ── Totals ───────────────────────────────────────────────
    building_subtotal = sum(r["original_budget"] for r in all_rows if r["div"] != "02")
    site_subtotal = sum(r["original_budget"] for r in all_rows if r["div"] == "02")

    # GC markup: user-confirmed (pq) > rate_profile > project-type defaults
    _gc = rate_profile.get('gc_markups', {})
    _pt_json = pq.get('project_type', 'unknown')
    _type_gc_defaults_json = {
        'single_family':  (0.05, 0.06, 0.05, 0.025),
        'multi_family':   (0.08, 0.10, 0.08, 0.025),
        'mixed_use':      (0.08, 0.10, 0.08, 0.025),
        'commercial':     (0.10, 0.12, 0.10, 0.025),
    }
    _td_json = _type_gc_defaults_json.get(_pt_json, (0.10, 0.12, 0.10, 0.025))
    gc_cond_pct = pq.get('gc_conditions_pct') or _gc.get('conditions_pct') or _td_json[0]
    gc_prof_pct = pq.get('gc_profit_pct') or _gc.get('profit_pct') or _td_json[1]
    cont_pct = pq.get('gc_contingency_pct') or _gc.get('contingency_pct') or _td_json[2]
    perm_pct = pq.get('gc_permits_pct') or _gc.get('permits_bond_pct') or _td_json[3]

    gc_cond = round(building_subtotal * gc_cond_pct)
    gc_prof = round(building_subtotal * gc_prof_pct)
    cont = round(building_subtotal * cont_pct)
    perm = round(building_subtotal * perm_pct)
    building_total = building_subtotal + gc_cond + gc_prof + cont + perm

    site_fee = round(site_subtotal * 0.05)
    site_total = site_subtotal + site_fee
    project_total = building_total + site_total

    print(f"[JSON TOTALS] building_sub={building_subtotal:,.0f} site_sub={site_subtotal:,.0f} "
          f"GC={gc_cond+gc_prof+cont+perm:,.0f} ({gc_cond_pct+gc_prof_pct+cont_pct+perm_pct:.1%}) "
          f"project_total={project_total:,.0f}", flush=True)

    totals = {
        "building_subtotal": building_subtotal,
        "gc_conditions": gc_cond, "gc_conditions_pct": gc_cond_pct,
        "gc_profit": gc_prof, "gc_profit_pct": gc_prof_pct,
        "contingency": cont, "contingency_pct": cont_pct,
        "permits": perm, "permits_pct": perm_pct,
        "building_total": building_total,
        "site_subtotal": site_subtotal, "site_fee": site_fee,
        "site_total": site_total,
        "project_total": project_total,
        "low_estimate": round(project_total * 0.85),
        "high_estimate": round(project_total * 1.15),
    }

    # ── Scope Details ────────────────────────────────────────
    scope_details = []
    for row_data in consolidated:
        scope_details.append({
            "div": row_data["csi_division"],
            "trade": row_data["trade"],
            "scope_items": row_data["scope"].split("\n"),
            "sources": row_data["source_pages"],
            "quantities": row_data.get("quantities"),
        })

    # ── Flags ────────────────────────────────────────────────
    flags = []
    for page_ref in (failed_pages or []):
        flags.append({"severity": "HIGH", "flag": "Page parse failure",
                      "detail": f"{page_ref} failed after 4 attempts.", "source": page_ref})
    for finding in (addenda_findings or []):
        flags.append(finding)
    for vf in pq.get('_validation_findings', []):
        flags.append({"severity": "MEDIUM", "flag": "Quantity validation",
                      "detail": vf, "source": "Project quantities"})

    # ── Scalars ──────────────────────────────────────────────
    confidence = pq.get('_confidence', {})
    scalars = []
    for key, label in [
        ('project_name', 'Project Name'), ('project_address', 'Project Address'),
        ('project_type', 'Project Type'), ('construction_type', 'Construction Type'),
        ('unit_count', 'Unit Count'), ('floor_count', 'Floor Count'),
        ('total_building_sf', 'Total Building SF'), ('footprint_sf', 'Footprint SF'),
        ('perimeter_lf', 'Perimeter LF'), ('roof_type', 'Roof Type'),
    ]:
        scalars.append({"key": key, "label": label,
                        "value": pq.get(key, ""), "confidence": confidence.get(key, 'low')})

    return {
        "divisions": divisions,
        "totals": totals,
        "scope_details": scope_details,
        "flags": flags,
        "scalars": scalars,
        "pricing_source": "rate_tables" if rate_profile else "material_db",
    }
